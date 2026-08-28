"""
Invoice Generator GUI Application
================================

This script provides a desktop GUI for creating and tracking service invoices.  It is
designed to replace a spreadsheet‑based invoicing system with a more automated
approach while still keeping data files easy to inspect or edit.  The core features are:

* Persistent invoice numbering.  The next invoice number is stored in a simple
  JSON file (`settings.json`) so each new invoice is automatically assigned
  the correct sequential number.  If the file does not exist the first invoice
  will start at 1.
* Customer and invoice metadata entry (date, due date, client name/address,
  optional notes).
* A table for line items where each row contains a description, quantity, unit
  price, taxable flag, subtotal, GST and total.  You can add, edit (double‑click)
  or remove items via dedicated controls.
* Optional service catalogue support.  A CSV file (`service_items.csv`) can
  contain commonly used items with their price and tax status.  Selecting an
  item from the catalogue will automatically populate the description, price
  and taxable flag for a new line item.  The catalogue can also be managed
  from within the application via the Tools menu.
* Invoice totals (subtotal, GST, total) automatically update whenever items
  change.
* Saving invoices: the data for each invoice is appended to an
  `invoices.csv` file for record keeping.  A PDF version of the invoice is
  generated using the reportlab library and stored in the `invoices` folder.
  The PDF is automatically opened after saving.
* An Invoice History tab lets you browse and re‑open previously saved PDFs.
* A Settings dialog lets you configure your business name, ABN, address,
  payment details and GST rate.  This information is embedded in every PDF.
* Clearing the form after an invoice is saved so you can begin a new one
  immediately.

The application uses Tkinter for the GUI.  ReportLab is required to generate
PDFs; it is available in this environment.  If you run this script on a
different machine and ReportLab is not installed you can install it via
`pip install reportlab`.

To run the application simply execute this script with Python 3:

    python invoice_gui.py

The first time you run the script it will create a few auxiliary files and
folders in the working directory:

* `settings.json` – stores the next invoice number and business configuration.
* `service_items.csv` – a sample service catalogue; edit it to reflect your
  own commonly used services.
* `invoices/` – a folder where generated PDF invoices are saved.
* `invoices.csv` – a CSV log of all invoices generated through this app.

Feel free to customise the PDF output by editing the `create_pdf` function.
"""

import csv
import json
import logging
import os
import subprocess
import sys
import traceback
from datetime import datetime, timedelta
from pathlib import Path

import tkinter as tk
from tkinter import ttk, messagebox, filedialog

from app_log import setup_logging, get_logger, log_path
from app_theme import apply_theme, configure_tags, ROW, ROW_FG, LABEL_MUTED, LABEL_SUCCESS, LABEL_DANGER
from data_store import DataStore, INVOICE_FIELDS, INVOICE_STATUSES
from ledger_tab import LedgerTab
from reports_tab import ReportsTab
from smart_clipboard import bind_treeview_clipboard
from date_utils import DateEntry, fmt_display, storage_to_display, display_to_storage, fmt_or_blank, parse_date

_log = get_logger('gui')

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.platypus import (SimpleDocTemplate, Table, TableStyle,
                                Paragraph, Spacer, HRFlowable)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_RIGHT, TA_CENTER


# ---------------------------------------------------------------------------
# Default settings written on first run
# ---------------------------------------------------------------------------
_DEFAULT_SETTINGS = {
    'next_invoice_number': 1,
    'business_name': 'Your Business Name',
    'business_abn': '',
    'business_address': '',
    'business_phone': '',
    'business_email': '',
    'bank_name': '',
    'bank_bsb': '',
    'bank_account': '',
    'bank_account_name': '',
    'gst_rate': 0.10,
    'payment_terms_days': 30,
    'currency_symbol': '$',
    'thank_you_note': 'Thank you for your business!',
    'show_gst_not_registered': False,
    'pdf_save_mode': 'auto',
    'pdf_save_dir': '',
    'training_manager': '',
    'auto_backup_enabled': False,
    'backup_frequency':    'daily',
    'backup_dir':          '',
    'backup_keep':         10,
    'backup_on_exit':      True,
}

# ---------------------------------------------------------------------------
# Config tab inside SettingsDialog — added below in the class


_HARDCODED_VERSION = '1.32'   # updated each release

def _read_version() -> str:
    """Read version from VERSION file; fall back to hardcoded constant."""
    import sys
    candidates = [
        Path(__file__).parent / 'VERSION',          # dev / source tree
        Path(sys.executable).parent / 'VERSION',    # PyInstaller onefile next to exe
    ]
    for vf in candidates:
        try:
            text = vf.read_text(encoding='utf-8').strip()
            if text:
                return text
        except Exception:
            pass
    return _HARDCODED_VERSION


APP_VERSION = _read_version()


class AboutDialog(tk.Toplevel):
    """Simple About dialog showing version and app info."""

    def __init__(self, parent):
        super().__init__(parent)
        self.title('About Invoice Generator')
        self.resizable(False, False)
        self.grab_set()
        self._build()
        self.wait_window(self)

    def _build(self):
        pad = {'padx': 20, 'pady': 6}
        ttk.Label(self, text='Invoice Generator',
                  font=('TkDefaultFont', 14, 'bold')).pack(**pad)
        ttk.Label(self, text=f'Version {APP_VERSION}',
                  font=('TkDefaultFont', 11)).pack(padx=20, pady=2)
        ttk.Separator(self, orient='horizontal').pack(fill='x', padx=16, pady=8)
        info = [
            ('Data format', 'CSV — human-readable, portable'),
            ('PDF engine',  'ReportLab'),
            ('GUI toolkit', f'Tkinter / Python {sys.version.split()[0]}'),
        ]
        for label, value in info:
            row = ttk.Frame(self)
            row.pack(fill='x', padx=20, pady=1)
            ttk.Label(row, text=label + ':', width=14, anchor='e',
                      foreground='grey').pack(side='left')
            ttk.Label(row, text=value).pack(side='left', padx=6)
        ttk.Separator(self, orient='horizontal').pack(fill='x', padx=16, pady=8)
        ttk.Button(self, text='Close', command=self.destroy).pack(pady=(0, 12))


def _open_file(path: Path):
    """Open a file with the default system application."""
    try:
        if sys.platform == 'win32':
            os.startfile(str(path))
        elif sys.platform == 'darwin':
            subprocess.Popen(['open', str(path)])
        else:
            subprocess.Popen(['xdg-open', str(path)])
    except Exception as e:
        messagebox.showerror("Open Error", f"Could not open file:\n{e}")


def _detect_onedrive() -> str:
    """
    Return the user's OneDrive root folder path, or '' if not found.
    Checks (in order):
      1. ONEDRIVE  environment variable (set by the OneDrive client)
      2. OneDriveConsumer  environment variable
      3. USERPROFILE\\OneDrive  (default install location)
      4. Windows registry  HKCU\\Software\\Microsoft\\OneDrive  UserFolder
    """
    for env in ('OneDrive', 'OneDriveConsumer', 'ONEDRIVE'):
        val = os.environ.get(env, '').strip()
        if val and Path(val).is_dir():
            return val
    # Fallback: default location under user profile
    profile = os.environ.get('USERPROFILE', '')
    if profile:
        candidate = Path(profile) / 'OneDrive'
        if candidate.is_dir():
            return str(candidate)
    # Windows registry (best effort)
    try:
        import winreg
        key = winreg.OpenKey(winreg.HKEY_CURRENT_USER,
                             r'Software\Microsoft\OneDrive')
        val, _ = winreg.QueryValueEx(key, 'UserFolder')
        winreg.CloseKey(key)
        if val and Path(val).is_dir():
            return val
    except Exception:
        pass
    return ''


class SettingsDialog(tk.Toplevel):
    """Modal dialog for editing business and payment settings."""

    def __init__(self, parent, settings: dict, backup_mgr=None):
        super().__init__(parent)
        self.title("Settings")
        self.resizable(False, False)
        self.grab_set()
        self.settings = dict(settings)
        self.result = None
        self._backup_mgr_ref = backup_mgr
        self._build()
        self.wait_window(self)

    def _build(self):
        pad = {'padx': 6, 'pady': 3}

        nb = ttk.Notebook(self)
        nb.pack(fill='both', expand=True, padx=10, pady=10)

        # ---- Business tab ----
        biz = ttk.Frame(nb)
        nb.add(biz, text='Business')

        fields_biz = [
            ('Business name:', 'business_name', 40),
            ('ABN:', 'business_abn', 20),
            ('Address:', 'business_address', 40),
            ('Phone:', 'business_phone', 20),
            ('Email:', 'business_email', 30),
        ]
        self._biz_vars = {}
        for r, (label, key, width) in enumerate(fields_biz):
            ttk.Label(biz, text=label).grid(row=r, column=0, sticky='e', **pad)
            var = tk.StringVar(value=self.settings.get(key, ''))
            self._biz_vars[key] = var
            ttk.Entry(biz, textvariable=var, width=width).grid(row=r, column=1, sticky='w', **pad)

        # ---- Payment tab ----
        pay = ttk.Frame(nb)
        nb.add(pay, text='Payment')

        fields_pay = [
            ('Bank name:', 'bank_name', 30),
            ('BSB:', 'bank_bsb', 15),
            ('Account number:', 'bank_account', 20),
            ('Account name:', 'bank_account_name', 30),
        ]
        self._pay_vars = {}
        for r, (label, key, width) in enumerate(fields_pay):
            ttk.Label(pay, text=label).grid(row=r, column=0, sticky='e', **pad)
            var = tk.StringVar(value=self.settings.get(key, ''))
            self._pay_vars[key] = var
            ttk.Entry(pay, textvariable=var, width=width).grid(row=r, column=1, sticky='w', **pad)

        # ---- Invoice tab ----
        inv = ttk.Frame(nb)
        nb.add(inv, text='Invoice')

        ttk.Label(inv, text="Next invoice number:").grid(row=0, column=0, sticky='e', **pad)
        self._next_inv_var = tk.StringVar(value=str(self.settings.get('next_invoice_number', 1)))
        ttk.Entry(inv, textvariable=self._next_inv_var, width=8).grid(row=0, column=1, sticky='w', **pad)
        ttk.Label(inv, text='The number that will be used on the next invoice.',
                  foreground='grey').grid(row=0, column=2, sticky='w', padx=4)

        ttk.Label(inv, text="GST rate (%):").grid(row=1, column=0, sticky='e', **pad)
        self._gst_var = tk.StringVar(value=f"{self.settings.get('gst_rate', 0.10) * 100:.1f}")
        ttk.Entry(inv, textvariable=self._gst_var, width=8).grid(row=1, column=1, sticky='w', **pad)

        ttk.Label(inv, text="Payment terms (days):").grid(row=2, column=0, sticky='e', **pad)
        self._terms_var = tk.StringVar(value=str(self.settings.get('payment_terms_days', 30)))
        ttk.Entry(inv, textvariable=self._terms_var, width=8).grid(row=2, column=1, sticky='w', **pad)

        ttk.Label(inv, text="Currency symbol:").grid(row=3, column=0, sticky='e', **pad)
        self._currency_var = tk.StringVar(value=self.settings.get('currency_symbol', '$'))
        ttk.Entry(inv, textvariable=self._currency_var, width=5).grid(row=3, column=1, sticky='w', **pad)

        # ---- PDF tab ----
        pdf_tab = ttk.Frame(nb)
        nb.add(pdf_tab, text='PDF')

        ttk.Label(pdf_tab, text="Training Manager name:").grid(row=0, column=0, sticky='e', **pad)
        self._training_mgr_var = tk.StringVar(value=self.settings.get('training_manager', ''))
        ttk.Entry(pdf_tab, textvariable=self._training_mgr_var, width=35).grid(row=0, column=1, sticky='w', **pad)
        ttk.Label(pdf_tab, text='Used on course completion certificates.',
                  foreground='grey').grid(row=1, column=1, sticky='w', padx=6)

        ttk.Label(pdf_tab, text="Thank-you note:").grid(row=2, column=0, sticky='ne', **pad)
        self._thankyou_var = tk.StringVar(value=self.settings.get('thank_you_note', 'Thank you for your business!'))
        ttk.Entry(pdf_tab, textvariable=self._thankyou_var, width=45).grid(row=2, column=1, sticky='w', **pad)
        ttk.Label(pdf_tab, text="Leave blank to omit.", foreground='grey').grid(row=3, column=1, sticky='w', padx=6)

        self._gst_note_var = tk.BooleanVar(value=self.settings.get('show_gst_not_registered', False))
        ttk.Checkbutton(pdf_tab, text='Show "Not registered for GST" note on invoice',
                        variable=self._gst_note_var).grid(row=4, column=0, columnspan=2, sticky='w', padx=6, pady=6)

        ttk.Separator(pdf_tab, orient='horizontal').grid(row=5, column=0, columnspan=3, sticky='ew', padx=6, pady=8)
        ttk.Label(pdf_tab, text="PDF save location:", font=('TkDefaultFont', 9, 'bold')).grid(
            row=6, column=0, columnspan=2, sticky='w', padx=6)

        self._pdf_mode_var = tk.StringVar(value=self.settings.get('pdf_save_mode', 'auto'))
        ttk.Radiobutton(pdf_tab, text='Auto – save to default folder (no dialog)',
                        variable=self._pdf_mode_var, value='auto').grid(
            row=7, column=0, columnspan=2, sticky='w', padx=6, pady=2)
        ttk.Radiobutton(pdf_tab, text='Prompt – ask where to save each time',
                        variable=self._pdf_mode_var, value='prompt').grid(
            row=8, column=0, columnspan=2, sticky='w', padx=6, pady=2)

        dir_frame = ttk.Frame(pdf_tab)
        dir_frame.grid(row=9, column=0, columnspan=2, sticky='ew', padx=6, pady=4)
        ttk.Label(dir_frame, text="Default folder:").pack(side='left')
        self._pdf_dir_var = tk.StringVar(value=self.settings.get('pdf_save_dir', ''))
        ttk.Entry(dir_frame, textvariable=self._pdf_dir_var, width=32).pack(side='left', padx=4)
        ttk.Button(dir_frame, text="Browse…",
                   command=self._browse_pdf_dir).pack(side='left')
        ttk.Label(pdf_tab, text="Leave blank to use the invoices/ sub-folder.",
                  foreground='grey').grid(row=10, column=0, columnspan=2, sticky='w', padx=6)

        # ---- Config tab ----
        cfg_tab = ttk.Frame(nb)
        nb.add(cfg_tab, text='Config')
        pad2 = {'padx': 8, 'pady': 4}

        ttk.Label(cfg_tab, text='Data directory:', font=('TkDefaultFont', 9, 'bold')).grid(
            row=0, column=0, columnspan=3, sticky='w', **pad2)
        ttk.Label(cfg_tab, text='All CSV data files are stored here.',
                  foreground='grey').grid(row=1, column=0, columnspan=3, sticky='w', padx=8)
        dir_row = ttk.Frame(cfg_tab)
        dir_row.grid(row=2, column=0, columnspan=3, sticky='ew', **pad2)
        self._data_dir_var = tk.StringVar(value=self.settings.get('_data_dir', ''))
        ttk.Entry(dir_row, textvariable=self._data_dir_var, width=42).pack(side='left', padx=4)
        ttk.Button(dir_row, text='Browse…', command=self._browse_data_dir).pack(side='left')
        ttk.Label(cfg_tab, text='Leave blank to use the same folder as the application.',
                  foreground='grey').grid(row=3, column=0, columnspan=3, sticky='w', padx=8)

        # OneDrive quick-setup row
        od_frame = ttk.LabelFrame(cfg_tab, text='OneDrive')
        od_frame.grid(row=4, column=0, columnspan=3, sticky='ew', padx=8, pady=(4, 0))
        od_path = _detect_onedrive()
        od_suggestion = str(Path(od_path) / 'InvoicerData') if od_path else ''
        if od_path:
            od_status = f'Detected: {od_path}'
            od_colour = '#155724'
        else:
            od_status = 'OneDrive folder not detected on this machine.'
            od_colour = '#856404'
        ttk.Label(od_frame, text=od_status, foreground=od_colour).grid(
            row=0, column=0, columnspan=3, sticky='w', padx=8, pady=2)
        ttk.Button(od_frame, text='Use OneDrive folder',
                   command=lambda: self._data_dir_var.set(od_suggestion) if od_suggestion
                           else messagebox.showinfo('Not found',
                               'OneDrive folder could not be detected.\n'
                               'Browse manually to your OneDrive folder.', parent=self)
                   ).grid(row=1, column=0, sticky='w', padx=8, pady=4)
        ttk.Label(od_frame,
                  text=f'Sets data folder to: {od_suggestion}' if od_suggestion
                       else 'Browse to your OneDrive folder manually.',
                  foreground='grey').grid(row=1, column=1, sticky='w', padx=4)
        ttk.Button(od_frame, text='Move existing data there…',
                   command=self._move_data_to_onedrive).grid(
                   row=2, column=0, sticky='w', padx=8, pady=(0, 6))
        ttk.Label(od_frame, text='Copies all current data files to the folder above, then switches.',
                  foreground='grey').grid(row=2, column=1, sticky='w', padx=4)

        ttk.Separator(cfg_tab, orient='horizontal').grid(row=5, column=0, columnspan=3,
                                                         sticky='ew', padx=6, pady=8)
        ttk.Label(cfg_tab, text='Invoice PDF save folder:', font=('TkDefaultFont', 9, 'bold')).grid(
            row=6, column=0, columnspan=3, sticky='w', **pad2)
        inv_dir_row = ttk.Frame(cfg_tab)
        inv_dir_row.grid(row=7, column=0, columnspan=3, sticky='ew', **pad2)
        self._inv_save_dir_var = tk.StringVar(value=self.settings.get('pdf_save_dir', ''))
        ttk.Entry(inv_dir_row, textvariable=self._inv_save_dir_var, width=40).pack(side='left', padx=4)
        ttk.Button(inv_dir_row, text='Browse…', command=self._browse_inv_dir).pack(side='left')
        ttk.Label(cfg_tab, text='Leave blank to use invoices/ subfolder.',
                  foreground='grey').grid(row=8, column=0, columnspan=3, sticky='w', padx=8)

        # ---- Reports tab ----
        rpt_tab = ttk.Frame(nb)
        nb.add(rpt_tab, text='Reports')
        pad3 = {'padx': 8, 'pady': 4}

        ttk.Label(rpt_tab, text='PDF Colour Scheme',
                  font=('TkDefaultFont', 9, 'bold')).grid(
            row=0, column=0, columnspan=3, sticky='w', **pad3)
        ttk.Label(rpt_tab, text='Header colour (hex):').grid(row=1, column=0, sticky='e', **pad3)
        self._rpt_header_colour_var = tk.StringVar(
            value=self.settings.get('report_header_colour', '#2C3E50'))
        ttk.Entry(rpt_tab, textvariable=self._rpt_header_colour_var, width=10).grid(
            row=1, column=1, sticky='w', **pad3)
        ttk.Label(rpt_tab, text='e.g. #2C3E50 (dark navy)',
                  foreground='grey').grid(row=1, column=2, sticky='w', padx=4)

        ttk.Label(rpt_tab, text='Accent colour (hex):').grid(row=2, column=0, sticky='e', **pad3)
        self._rpt_accent_colour_var = tk.StringVar(
            value=self.settings.get('report_accent_colour', '#2980B9'))
        ttk.Entry(rpt_tab, textvariable=self._rpt_accent_colour_var, width=10).grid(
            row=2, column=1, sticky='w', **pad3)
        ttk.Label(rpt_tab, text='e.g. #2980B9 (blue) — used for column headers',
                  foreground='grey').grid(row=2, column=2, sticky='w', padx=4)

        ttk.Label(rpt_tab, text='Row stripe colour (hex):').grid(row=3, column=0, sticky='e', **pad3)
        self._rpt_stripe_colour_var = tk.StringVar(
            value=self.settings.get('report_stripe_colour', '#EBF5FB'))
        ttk.Entry(rpt_tab, textvariable=self._rpt_stripe_colour_var, width=10).grid(
            row=3, column=1, sticky='w', **pad3)
        ttk.Label(rpt_tab, text='e.g. #EBF5FB — alternating row tint',
                  foreground='grey').grid(row=3, column=2, sticky='w', padx=4)

        ttk.Separator(rpt_tab, orient='horizontal').grid(
            row=4, column=0, columnspan=3, sticky='ew', padx=6, pady=8)
        ttk.Label(rpt_tab, text='Wording',
                  font=('TkDefaultFont', 9, 'bold')).grid(
            row=5, column=0, columnspan=3, sticky='w', **pad3)

        ttk.Label(rpt_tab, text='Report prepared by:').grid(row=6, column=0, sticky='e', **pad3)
        self._rpt_prepared_by_var = tk.StringVar(
            value=self.settings.get('report_prepared_by', ''))
        ttk.Entry(rpt_tab, textvariable=self._rpt_prepared_by_var, width=32).grid(
            row=6, column=1, columnspan=2, sticky='w', **pad3)

        ttk.Label(rpt_tab, text='Report footer note:').grid(row=7, column=0, sticky='e', **pad3)
        self._rpt_footer_var = tk.StringVar(
            value=self.settings.get('report_footer', 'CONFIDENTIAL — For internal use only.'))
        ttk.Entry(rpt_tab, textvariable=self._rpt_footer_var, width=45).grid(
            row=7, column=1, columnspan=2, sticky='w', **pad3)

        ttk.Label(rpt_tab, text='Organisation name override:').grid(row=8, column=0, sticky='e', **pad3)
        self._rpt_org_var = tk.StringVar(
            value=self.settings.get('report_org_override', ''))
        ttk.Entry(rpt_tab, textvariable=self._rpt_org_var, width=45).grid(
            row=8, column=1, columnspan=2, sticky='w', **pad3)
        ttk.Label(rpt_tab, text='Leave blank to use business name from Business tab.',
                  foreground='grey').grid(row=9, column=1, columnspan=2, sticky='w', padx=6)

        # ---- Backup tab ----
        bkp_tab = ttk.Frame(nb)
        nb.add(bkp_tab, text='Backup')
        pad4 = {'padx': 8, 'pady': 4}

        ttk.Label(bkp_tab, text='Automatic Backups',
                  font=('TkDefaultFont', 9, 'bold')).grid(
            row=0, column=0, columnspan=3, sticky='w', **pad4)
        ttk.Label(bkp_tab,
                  text='Backups are written silently as timestamped zips to the chosen folder.',
                  foreground='grey').grid(row=1, column=0, columnspan=3, sticky='w', padx=8)

        self._bkp_enabled_var = tk.BooleanVar(value=self.settings.get('auto_backup_enabled', False))
        ttk.Checkbutton(bkp_tab, text='Enable automatic backups',
                        variable=self._bkp_enabled_var).grid(
            row=2, column=0, columnspan=2, sticky='w', padx=8, pady=(8, 2))

        ttk.Label(bkp_tab, text='Frequency:').grid(row=3, column=0, sticky='e', **pad4)
        self._bkp_freq_var = tk.StringVar(value=self.settings.get('backup_frequency', 'daily'))
        freq_cb = ttk.Combobox(bkp_tab, textvariable=self._bkp_freq_var,
                               values=['hourly', 'every2h', 'every4h',
                                       'every6h', 'every12h', 'daily'],
                               state='readonly', width=14)
        freq_cb.grid(row=3, column=1, sticky='w', **pad4)
        ttk.Label(bkp_tab,
                  text='hourly / every2h / every4h / every6h / every12h / daily',
                  foreground='grey').grid(row=3, column=2, sticky='w', padx=4)

        ttk.Label(bkp_tab, text='Backup folder:').grid(row=4, column=0, sticky='e', **pad4)
        bkp_dir_row = ttk.Frame(bkp_tab)
        bkp_dir_row.grid(row=4, column=1, columnspan=2, sticky='ew', **pad4)
        self._bkp_dir_var = tk.StringVar(value=self.settings.get('backup_dir', ''))
        ttk.Entry(bkp_dir_row, textvariable=self._bkp_dir_var, width=38).pack(side='left', padx=4)
        ttk.Button(bkp_dir_row, text='Browse…',
                   command=self._browse_backup_dir).pack(side='left')
        ttk.Label(bkp_tab,
                  text='Leave blank to use a backups/ subfolder inside the data directory.',
                  foreground='grey').grid(row=5, column=1, columnspan=2, sticky='w', padx=8)

        ttk.Label(bkp_tab, text='Keep last N copies:').grid(row=6, column=0, sticky='e', **pad4)
        self._bkp_keep_var = tk.StringVar(value=str(self.settings.get('backup_keep', 10)))
        ttk.Entry(bkp_tab, textvariable=self._bkp_keep_var, width=6).grid(
            row=6, column=1, sticky='w', **pad4)
        ttk.Label(bkp_tab, text='Older backups are deleted automatically.',
                  foreground='grey').grid(row=6, column=2, sticky='w', padx=4)

        self._bkp_on_exit_var = tk.BooleanVar(value=self.settings.get('backup_on_exit', True))
        ttk.Checkbutton(bkp_tab, text='Also backup when the application closes',
                        variable=self._bkp_on_exit_var).grid(
            row=7, column=0, columnspan=3, sticky='w', padx=8, pady=(8, 2))

        ttk.Separator(bkp_tab, orient='horizontal').grid(
            row=8, column=0, columnspan=3, sticky='ew', padx=6, pady=8)

        ttk.Label(bkp_tab, text='Existing backups',
                  font=('TkDefaultFont', 9, 'bold')).grid(
            row=9, column=0, columnspan=3, sticky='w', **pad4)
        self._bkp_status_var = tk.StringVar(value='(click Refresh to list)')
        ttk.Label(bkp_tab, textvariable=self._bkp_status_var,
                  foreground='#555').grid(row=10, column=0, columnspan=3, sticky='w', padx=8)
        bkp_btn_row = ttk.Frame(bkp_tab)
        bkp_btn_row.grid(row=11, column=0, columnspan=3, sticky='w', padx=8, pady=4)
        ttk.Button(bkp_btn_row, text='Backup Now',
                   command=self._bkp_run_now).pack(side='left', padx=3)
        ttk.Button(bkp_btn_row, text='Refresh list',
                   command=self._bkp_refresh_list).pack(side='left', padx=3)
        ttk.Button(bkp_btn_row, text='Open folder',
                   command=self._bkp_open_folder).pack(side='left', padx=3)

        self._bkp_refresh_list()

        # ---- Startup tab ----
        start_tab = ttk.Frame(nb)
        nb.add(start_tab, text='Startup')
        pad5 = {'padx': 8, 'pady': 4}
        ttk.Label(start_tab, text='Windows Startup',
                  font=('TkDefaultFont', 9, 'bold')).grid(
            row=0, column=0, columnspan=2, sticky='w', **pad5)
        ttk.Label(start_tab,
                  text='Launch Invoice Generator automatically when you log in to Windows.',
                  foreground='grey').grid(row=1, column=0, columnspan=2, sticky='w', padx=8)
        import startup_manager
        self._startup_var = tk.BooleanVar(value=startup_manager.is_enabled())
        ttk.Checkbutton(start_tab, text='Start Invoice Generator on Windows startup (Delayed Run)',
                        variable=self._startup_var).grid(
            row=2, column=0, columnspan=2, sticky='w', padx=8, pady=(8, 2))
        ttk.Label(start_tab,
                  text='This entry also appears in Windows Settings → Apps → Startup,\n'
                       'where it can be toggled on or off.',
                  foreground='grey').grid(row=3, column=0, columnspan=2, sticky='w', padx=8, pady=(4, 0))

        # ---- Buttons ----
        btn_frame = ttk.Frame(self)
        btn_frame.pack(fill='x', padx=10, pady=(0, 10))
        ttk.Button(btn_frame, text="Save", command=self._on_save).pack(side='right', padx=5)
        ttk.Button(btn_frame, text="Cancel", command=self.destroy).pack(side='right')

    def _on_save(self):
        try:
            next_inv = int(self._next_inv_var.get())
            if next_inv < 1:
                raise ValueError
        except ValueError:
            messagebox.showwarning("Invalid invoice number",
                                   "Next invoice number must be a whole number ≥ 1.", parent=self)
            return
        try:
            gst_rate = float(self._gst_var.get()) / 100.0
        except ValueError:
            messagebox.showwarning("Invalid GST", "GST rate must be a number.", parent=self)
            return
        try:
            terms = int(self._terms_var.get())
        except ValueError:
            messagebox.showwarning("Invalid terms", "Payment terms must be a whole number.", parent=self)
            return
        self.result = dict(self.settings)
        for key, var in {**self._biz_vars, **self._pay_vars}.items():
            self.result[key] = var.get().strip()
        self.result['next_invoice_number'] = next_inv
        self.result['gst_rate'] = gst_rate
        self.result['payment_terms_days'] = terms
        self.result['currency_symbol'] = self._currency_var.get()
        self.result['thank_you_note'] = self._thankyou_var.get().strip()
        self.result['show_gst_not_registered'] = self._gst_note_var.get()
        self.result['training_manager'] = self._training_mgr_var.get().strip()
        self.result['pdf_save_mode'] = self._pdf_mode_var.get()
        self.result['pdf_save_dir'] = self._pdf_dir_var.get().strip()
        self.result['_data_dir'] = self._data_dir_var.get().strip()
        # inv_save_dir (Config tab) takes precedence over PDF tab entry when set
        inv_dir = self._inv_save_dir_var.get().strip()
        if inv_dir:
            self.result['pdf_save_dir'] = inv_dir
        self.result['report_header_colour'] = self._rpt_header_colour_var.get().strip() or '#2C3E50'
        self.result['report_accent_colour']  = self._rpt_accent_colour_var.get().strip() or '#2980B9'
        self.result['report_stripe_colour']  = self._rpt_stripe_colour_var.get().strip() or '#EBF5FB'
        self.result['report_prepared_by']    = self._rpt_prepared_by_var.get().strip()
        self.result['report_footer']         = self._rpt_footer_var.get().strip()
        self.result['report_org_override']   = self._rpt_org_var.get().strip()
        self.result['auto_backup_enabled'] = self._bkp_enabled_var.get()
        self.result['backup_frequency']    = self._bkp_freq_var.get()
        self.result['backup_dir']          = self._bkp_dir_var.get().strip()
        try:
            self.result['backup_keep'] = max(1, int(self._bkp_keep_var.get()))
        except (ValueError, TypeError):
            self.result['backup_keep'] = 10
        self.result['backup_on_exit'] = self._bkp_on_exit_var.get()
        # Apply Windows startup preference
        try:
            import startup_manager
            startup_manager.set_enabled(self._startup_var.get())
            self.result['run_on_startup'] = self._startup_var.get()
        except Exception:
            pass
        self.destroy()

    def _browse_backup_dir(self):
        from tkinter import filedialog as _fd
        chosen = _fd.askdirectory(title='Choose backup folder',
                                  initialdir=self._bkp_dir_var.get() or None,
                                  parent=self)
        if chosen:
            self._bkp_dir_var.set(chosen)

    def _bkp_refresh_list(self):
        from auto_backup import AutoBackupManager
        from pathlib import Path as _P
        bdir_str = self._bkp_dir_var.get().strip()
        bdir = _P(bdir_str) if bdir_str else None
        if bdir is None:
            self._bkp_status_var.set('(save settings first to resolve backup folder)')
            return
        backups = AutoBackupManager.list_backups(bdir)
        if not backups:
            self._bkp_status_var.set('No backups found in that folder yet.')
        else:
            lines = [f"{b['modified']}  {b['name']}  ({b['size_kb']} KB)"
                     for b in backups[:8]]
            if len(backups) > 8:
                lines.append(f'  … and {len(backups)-8} more')
            self._bkp_status_var.set('\n'.join(lines))

    def _bkp_run_now(self):
        if hasattr(self, '_backup_mgr_ref') and self._backup_mgr_ref:
            self._backup_mgr_ref.run_now(on_exit=False)
            self._bkp_refresh_list()
            messagebox.showinfo('Backup complete', 'Backup written successfully.', parent=self)
        else:
            messagebox.showinfo('Not available',
                'Open Settings from within the running app to use this button.', parent=self)

    def _bkp_open_folder(self):
        from pathlib import Path as _P
        import subprocess as _sp
        bdir_str = self._bkp_dir_var.get().strip()
        if not bdir_str:
            messagebox.showinfo('No folder', 'Set and save the backup folder first.', parent=self)
            return
        _sp.Popen(['explorer', str(_P(bdir_str))], shell=True)

    def _browse_pdf_dir(self):
        from tkinter import filedialog as _fd
        chosen = _fd.askdirectory(title="Choose default PDF save folder",
                                  initialdir=self._pdf_dir_var.get() or None,
                                  parent=self)
        if chosen:
            self._pdf_dir_var.set(chosen)

    def _move_data_to_onedrive(self):
        """Copy all current data files to the folder shown in the data-dir entry."""
        dest_str = self._data_dir_var.get().strip()
        if not dest_str:
            messagebox.showwarning('No destination',
                'Enter or choose a destination folder first (e.g. click "Use OneDrive folder").',
                parent=self)
            return
        dest = Path(dest_str)
        # Determine current data dir from the settings passed in
        current_str = self.settings.get('_data_dir', '').strip()
        if not current_str:
            messagebox.showinfo('Already set',
                'The data directory is currently the application folder.\n'
                'Click OK then Save — the app will start using the new folder.\n'
                'Your existing files will NOT be moved automatically; copy them manually if needed.',
                parent=self)
            return
        current = Path(current_str)
        if current.resolve() == dest.resolve():
            messagebox.showinfo('Same folder', 'Source and destination are the same folder.',
                parent=self)
            return
        if not messagebox.askyesno('Confirm move',
                f'Copy all data files from:\n  {current}\nto:\n  {dest}\n\n'
                'Original files will NOT be deleted — you can remove them manually after '
                'confirming everything works.',
                parent=self):
            return
        dest.mkdir(parents=True, exist_ok=True)
        import shutil as _sh
        copied, skipped = [], []
        for item in current.iterdir():
            if item.is_file():
                try:
                    _sh.copy2(item, dest / item.name)
                    copied.append(item.name)
                except Exception as e:
                    skipped.append(f'{item.name}: {e}')
            elif item.is_dir() and item.name == 'invoices':
                try:
                    _sh.copytree(item, dest / 'invoices', dirs_exist_ok=True)
                    copied.append('invoices/')
                except Exception as e:
                    skipped.append(f'invoices/: {e}')
        msg = f'Copied {len(copied)} item(s) to:\n  {dest}'
        if skipped:
            msg += '\n\nSkipped:\n' + '\n'.join(skipped)
        msg += '\n\nData folder has been updated to the new location.'
        self._data_dir_var.set(str(dest))
        messagebox.showinfo('Done', msg, parent=self)

    def _browse_data_dir(self):
        from tkinter import filedialog as _fd
        chosen = _fd.askdirectory(title='Choose data directory',
                                  initialdir=self._data_dir_var.get() or None,
                                  parent=self)
        if chosen:
            self._data_dir_var.set(chosen)

    def _browse_inv_dir(self):
        from tkinter import filedialog as _fd
        chosen = _fd.askdirectory(title='Choose invoice PDF save folder',
                                  initialdir=self._inv_save_dir_var.get() or None,
                                  parent=self)
        if chosen:
            self._inv_save_dir_var.set(chosen)


class CatalogueDialog(tk.Toplevel):
    """Modal dialog for managing service catalogue items."""

    def __init__(self, parent, service_items: list, save_path: Path):
        super().__init__(parent)
        self.title("Service Catalogue")
        self.minsize(500, 350)
        self.grab_set()
        self.service_items = [dict(i) for i in service_items]
        self.save_path = save_path
        self.changed = False
        self._build()
        self.wait_window(self)

    def _build(self):
        frame = ttk.Frame(self)
        frame.pack(fill='both', expand=True, padx=10, pady=10)

        columns = ('description', 'unit_price', 'taxable')
        self.tree = ttk.Treeview(frame, columns=columns, show='headings', height=10)
        self.tree.heading('description', text='Description')
        self.tree.heading('unit_price', text='Unit Price')
        self.tree.heading('taxable', text='Taxable')
        self.tree.column('description', width=280)
        self.tree.column('unit_price', width=90, anchor='e')
        self.tree.column('taxable', width=70, anchor='center')
        sb = ttk.Scrollbar(frame, orient='vertical', command=self.tree.yview)
        self.tree.configure(yscroll=sb.set)
        self.tree.grid(row=0, column=0, sticky='nsew')
        sb.grid(row=0, column=1, sticky='ns')
        frame.columnconfigure(0, weight=1)
        frame.rowconfigure(0, weight=1)

        self._refresh_tree()

        # Entry row
        entry_frame = ttk.LabelFrame(self, text="Add / Edit Item")
        entry_frame.pack(fill='x', padx=10, pady=5)

        ttk.Label(entry_frame, text="Description:").grid(row=0, column=0, sticky='e', padx=4, pady=3)
        self._desc_var = tk.StringVar()
        ttk.Entry(entry_frame, textvariable=self._desc_var, width=35).grid(row=0, column=1, sticky='w', padx=4)

        ttk.Label(entry_frame, text="Unit price:").grid(row=0, column=2, sticky='e', padx=4)
        self._price_var = tk.StringVar()
        ttk.Entry(entry_frame, textvariable=self._price_var, width=10).grid(row=0, column=3, sticky='w', padx=4)

        self._taxable_var = tk.BooleanVar()
        ttk.Checkbutton(entry_frame, text="Taxable", variable=self._taxable_var).grid(row=0, column=4, padx=4)

        ttk.Button(entry_frame, text="Add / Update", command=self._on_add).grid(row=0, column=5, padx=4)

        # Populate on double-click
        self.tree.bind('<Double-1>', self._on_double_click)

        btn_frame = ttk.Frame(self)
        btn_frame.pack(fill='x', padx=10, pady=(0, 10))
        ttk.Button(btn_frame, text="Remove Selected", command=self._on_remove).pack(side='left')
        ttk.Button(btn_frame, text="Save & Close", command=self._on_save).pack(side='right', padx=5)
        ttk.Button(btn_frame, text="Cancel", command=self.destroy).pack(side='right')

    def _refresh_tree(self):
        for row in self.tree.get_children():
            self.tree.delete(row)
        for item in self.service_items:
            self.tree.insert('', 'end', values=(
                item['description'],
                f"{item['unit_price']:.2f}",
                'Yes' if item['taxable'] else 'No'
            ))

    def _on_double_click(self, event=None):
        sel = self.tree.selection()
        if not sel:
            return
        idx = self.tree.index(sel[0])
        item = self.service_items[idx]
        self._desc_var.set(item['description'])
        self._price_var.set(f"{item['unit_price']:.2f}")
        self._taxable_var.set(item['taxable'])

    def _on_add(self):
        desc = self._desc_var.get().strip()
        if not desc:
            messagebox.showwarning("Missing description", "Enter a description.", parent=self)
            return
        try:
            price = float(self._price_var.get())
        except ValueError:
            messagebox.showwarning("Invalid price", "Unit price must be a number.", parent=self)
            return
        taxable = self._taxable_var.get()
        # Update existing or append
        for item in self.service_items:
            if item['description'].lower() == desc.lower():
                item['unit_price'] = price
                item['taxable'] = taxable
                self._refresh_tree()
                return
        self.service_items.append({'description': desc, 'unit_price': price, 'taxable': taxable})
        self._refresh_tree()

    def _on_remove(self):
        sel = self.tree.selection()
        if not sel:
            return
        idx = self.tree.index(sel[0])
        self.service_items.pop(idx)
        self._refresh_tree()

    def _on_save(self):
        with open(self.save_path, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            writer.writerow(['description', 'unit_price', 'taxable'])
            for item in self.service_items:
                writer.writerow([item['description'], f"{item['unit_price']:.2f}",
                                  'yes' if item['taxable'] else 'no'])
        self.changed = True
        self.destroy()


class ClientsDialog(tk.Toplevel):
    """Modal dialog for adding, editing and deleting clients."""

    FIELDS = [
        ('Name *',        'name',         35),
        ('Contact name',  'contact_name', 30),
        ('Phone',         'phone',        20),
        ('Email',         'email',        30),
        ('Address',       'address',      35),
    ]

    def __init__(self, parent, clients: list, save_path: Path):
        super().__init__(parent)
        self.title("Manage Clients")
        self.minsize(680, 420)
        self.grab_set()
        self.clients = [dict(c) for c in clients]
        self.save_path = save_path
        self.changed = False
        self._editing_idx = None
        self._build()
        self.wait_window(self)

    def _build(self):
        # ---- client list ----
        list_frame = ttk.Frame(self)
        list_frame.pack(fill='both', expand=True, padx=10, pady=(10, 4))

        cols = ('name', 'contact_name', 'phone', 'email')
        self.tree = ttk.Treeview(list_frame, columns=cols, show='headings',
                                 selectmode='browse', height=10)
        self.tree.heading('name',         text='Name')
        self.tree.heading('contact_name', text='Contact')
        self.tree.heading('phone',        text='Phone')
        self.tree.heading('email',        text='Email')
        self.tree.column('name',         width=170)
        self.tree.column('contact_name', width=130)
        self.tree.column('phone',        width=110)
        self.tree.column('email',        width=170)
        sb = ttk.Scrollbar(list_frame, orient='vertical', command=self.tree.yview)
        self.tree.configure(yscroll=sb.set)
        self.tree.grid(row=0, column=0, sticky='nsew')
        sb.grid(row=0, column=1, sticky='ns')
        list_frame.columnconfigure(0, weight=1)
        list_frame.rowconfigure(0, weight=1)
        self.tree.bind('<Double-1>', self._on_select)

        # ---- entry form ----
        entry_frame = ttk.LabelFrame(self, text="Client Details")
        entry_frame.pack(fill='x', padx=10, pady=4)

        self._vars = {}
        for r, (label, key, width) in enumerate(self.FIELDS):
            col_offset = (r % 2) * 2
            row_pos = r // 2
            ttk.Label(entry_frame, text=label + ':').grid(
                row=row_pos, column=col_offset, sticky='e', padx=4, pady=3)
            var = tk.StringVar()
            self._vars[key] = var
            if key == 'address':
                ttk.Entry(entry_frame, textvariable=var, width=width).grid(
                    row=row_pos, column=col_offset + 1, sticky='ew', padx=4,
                    columnspan=3)
            else:
                ttk.Entry(entry_frame, textvariable=var, width=width).grid(
                    row=row_pos, column=col_offset + 1, sticky='w', padx=4)

        entry_frame.columnconfigure(1, weight=1)
        entry_frame.columnconfigure(3, weight=1)

        # ---- buttons ----
        btn_frame = ttk.Frame(self)
        btn_frame.pack(fill='x', padx=10, pady=(4, 10))

        self._save_btn_text = tk.StringVar(value="Add Client")
        ttk.Button(btn_frame, textvariable=self._save_btn_text,
                   command=self._on_save_client).pack(side='left', padx=(0, 4))
        ttk.Button(btn_frame, text="Clear / New",
                   command=self._clear_form).pack(side='left', padx=4)
        ttk.Button(btn_frame, text="Remove Selected",
                   command=self._on_remove).pack(side='left', padx=4)
        ttk.Button(btn_frame, text="Save & Close",
                   command=self._on_done).pack(side='right', padx=(4, 0))
        ttk.Button(btn_frame, text="Cancel",
                   command=self.destroy).pack(side='right', padx=4)

        self._refresh_tree()

    # ------------------------------------------------------------------
    def _refresh_tree(self):
        for row in self.tree.get_children():
            self.tree.delete(row)
        for c in self.clients:
            self.tree.insert('', 'end', values=(
                c.get('name', ''),
                c.get('contact_name', ''),
                c.get('phone', ''),
                c.get('email', ''),
            ))

    def _clear_form(self):
        for var in self._vars.values():
            var.set('')
        self._editing_idx = None
        self._save_btn_text.set("Add Client")

    def _on_select(self, event=None):
        sel = self.tree.selection()
        if not sel:
            return
        idx = self.tree.index(sel[0])
        c = self.clients[idx]
        for key, var in self._vars.items():
            var.set(c.get(key, ''))
        self._editing_idx = idx
        self._save_btn_text.set("Update Client")

    def _on_save_client(self):
        name = self._vars['name'].get().strip()
        if not name:
            messagebox.showwarning("Missing name", "Client name is required.", parent=self)
            return
        record = {key: var.get().strip() for key, var in self._vars.items()}
        if self._editing_idx is not None:
            self.clients[self._editing_idx] = record
        else:
            # Prevent duplicate names
            existing = [c['name'].lower() for c in self.clients]
            if name.lower() in existing:
                messagebox.showwarning(
                    "Duplicate", f"A client named '{name}' already exists.", parent=self)
                return
            self.clients.append(record)
        self._refresh_tree()
        self._clear_form()

    def _on_remove(self):
        sel = self.tree.selection()
        if not sel:
            return
        idx = self.tree.index(sel[0])
        name = self.clients[idx]['name']
        if not messagebox.askyesno("Remove client",
                                   f"Remove '{name}'?\nThis does not delete past invoices.",
                                   parent=self):
            return
        self.clients.pop(idx)
        self._refresh_tree()
        self._clear_form()

    def _on_done(self):
        # Auto-save any unsaved entry in the form fields
        name = self._vars['name'].get().strip()
        if name:
            record = {key: var.get().strip() for key, var in self._vars.items()}
            existing_names = [c['name'].lower() for c in self.clients]
            if self._editing_idx is not None:
                self.clients[self._editing_idx] = record
            elif name.lower() not in existing_names:
                self.clients.append(record)
        with open(self.save_path, 'w', newline='', encoding='utf-8') as f:
            writer = csv.DictWriter(f, fieldnames=['name', 'contact_name', 'phone', 'email', 'address'])
            writer.writeheader()
            writer.writerows(self.clients)
        self.changed = True
        self.destroy()


class RecordMissingInvoiceDialog(tk.Toplevel):
    """
    Manually enter a past invoice that was never saved to the system
    (e.g. issued as a Word doc, sent by email, or created in another tool).
    Writes directly to invoices.csv — does NOT generate a PDF.
    """

    def __init__(self, parent, ds, settings, currency_fn, clients=None, prefill=None):
        super().__init__(parent)
        self._prefill = prefill or {}
        self.title('Edit Invoice' if prefill else 'Record Missing Invoice')
        self.resizable(False, False)
        self.grab_set()
        self.ds          = ds
        self.settings    = settings
        self.currency_fn = currency_fn
        self._clients    = clients or []
        self.result      = None
        self._build()
        self.wait_window(self)

    def _build(self):
        pad = {'padx': 8, 'pady': 4}
        sym = self.currency_fn()

        pf = self._prefill

        # Suggest next unused invoice number (ignored when prefilling)
        existing = {r.get('invoice_number', '') for r in self.ds.read_invoices()}
        next_n = self.settings.get('next_invoice_number', 1)
        while str(next_n) in existing or f'{next_n:04d}' in existing:
            next_n += 1
        suggested_num = pf.get('invoice_number', f'{next_n:04d}')

        info = ttk.LabelFrame(self, text='Invoice Details')
        info.pack(fill='x', padx=12, pady=(12, 4))

        text_fields = [
            ('Invoice #:',      'invoice_number', suggested_num,                    10),
            ('Client name:',    'client_name',    pf.get('client_name',    ''),     28),
            ('Client address:', 'client_address', pf.get('client_address', ''),     36),
        ]
        self._vars = {}
        row_offset = 0
        for i, (label, key, default, width) in enumerate(text_fields):
            ttk.Label(info, text=label).grid(row=i, column=0, sticky='e', **pad)
            var = tk.StringVar(value=default)
            ttk.Entry(info, textvariable=var, width=width).grid(
                row=i, column=1, sticky='w', **pad)
            self._vars[key] = var
            row_offset = i + 1

        # Date fields with DateEntry
        ttk.Label(info, text='Invoice date:').grid(row=row_offset, column=0, sticky='e', **pad)
        inv_date_val = storage_to_display(pf.get('invoice_date', '')) or fmt_display(datetime.now().date())
        self._vars['invoice_date'] = tk.StringVar(value=inv_date_val)
        DateEntry(info, textvariable=self._vars['invoice_date'], width=12).grid(
            row=row_offset, column=1, sticky='w', **pad)

        ttk.Label(info, text='Due date:').grid(row=row_offset + 1, column=0, sticky='e', **pad)
        due_date_val = storage_to_display(pf.get('due_date', '')) or ''
        self._vars['due_date'] = tk.StringVar(value=due_date_val)
        DateEntry(info, textvariable=self._vars['due_date'], width=12).grid(
            row=row_offset + 1, column=1, sticky='w', **pad)

        # Client autocomplete from known clients
        if self._clients:
            client_names = [c.get('name', '') for c in self._clients if c.get('name')]
            cb = ttk.Combobox(info, textvariable=self._vars['client_name'],
                              values=client_names, width=27)
            cb.grid(row=3, column=1, sticky='w', **pad)

        # Financial section
        fin = ttk.LabelFrame(self, text=f'Amounts ({sym})')
        fin.pack(fill='x', padx=12, pady=4)

        ttk.Label(fin, text='Subtotal:').grid(row=0, column=0, sticky='e', **pad)
        self._subtotal_var = tk.StringVar(value=pf.get('subtotal', ''))
        ttk.Entry(fin, textvariable=self._subtotal_var, width=12).grid(row=0, column=1, sticky='w', **pad)
        ttk.Label(fin, text='(before GST)', foreground='grey').grid(row=0, column=2, sticky='w')

        gst_rate = self.settings.get('gst_rate', 0.1)
        ttk.Label(fin, text=f'GST ({gst_rate*100:.0f}%):').grid(row=1, column=0, sticky='e', **pad)
        self._gst_var = tk.StringVar(value=pf.get('gst', ''))
        ttk.Entry(fin, textvariable=self._gst_var, width=12).grid(row=1, column=1, sticky='w', **pad)
        ttk.Button(fin, text='Auto-calc GST', command=self._auto_gst).grid(
            row=1, column=2, sticky='w', padx=4)

        ttk.Label(fin, text='Total:').grid(row=2, column=0, sticky='e', **pad)
        self._total_var = tk.StringVar(value=pf.get('total', ''))
        ttk.Entry(fin, textvariable=self._total_var, width=12).grid(row=2, column=1, sticky='w', **pad)
        ttk.Button(fin, text='Auto-calc Total', command=self._auto_total).grid(
            row=2, column=2, sticky='w', padx=4)

        # Payment status
        pay = ttk.LabelFrame(self, text='Payment')
        pay.pack(fill='x', padx=12, pady=4)

        self._paid_var = tk.BooleanVar(
            value=pf.get('paid', '').lower() in ('yes', 'true', '1'))
        ttk.Checkbutton(pay, text='Already paid', variable=self._paid_var).grid(
            row=0, column=0, columnspan=2, sticky='w', **pad)
        ttk.Label(pay, text='Date paid:').grid(row=1, column=0, sticky='e', **pad)
        self._paid_date_var = tk.StringVar(
            value=storage_to_display(pf.get('paid_date', '')) or '')
        DateEntry(pay, textvariable=self._paid_date_var, width=12).grid(row=1, column=1, sticky='w', **pad)
        ttk.Label(pay, text='Payment note:').grid(row=2, column=0, sticky='e', **pad)
        self._pay_note_var = tk.StringVar(value=pf.get('payment_note', ''))
        ttk.Entry(pay, textvariable=self._pay_note_var, width=36).grid(
            row=2, column=1, columnspan=2, sticky='w', **pad)

        # Invoice-level notes
        ttk.Label(self, text='Notes:').pack(anchor='w', padx=12)
        self._notes_text = tk.Text(self, width=52, height=3)
        self._notes_text.pack(fill='x', padx=12, pady=(0, 4))
        if pf.get('notes'):
            self._notes_text.insert('1.0', pf['notes'])

        btn_lbl = 'Save Changes' if self._prefill else 'Save Invoice'
        btn_frame = ttk.Frame(self)
        btn_frame.pack(fill='x', padx=12, pady=(4, 12))
        ttk.Button(btn_frame, text=btn_lbl, command=self._on_save).pack(side='right', padx=4)
        ttk.Button(btn_frame, text='Cancel', command=self.destroy).pack(side='right')
        ttk.Label(btn_frame, text='No PDF will be generated.',
                  foreground='grey').pack(side='left')

    def _auto_gst(self):
        try:
            sub = float(self._subtotal_var.get().replace(',', '').replace('$', ''))
            gst_rate = float(self.settings.get('gst_rate', 0.1))
            self._gst_var.set(f'{sub * gst_rate:.2f}')
            self._auto_total()
        except ValueError:
            pass

    def _auto_total(self):
        try:
            sub = float(self._subtotal_var.get().replace(',', '').replace('$', '') or 0)
            gst = float(self._gst_var.get().replace(',', '').replace('$', '') or 0)
            self._total_var.set(f'{sub + gst:.2f}')
        except ValueError:
            pass

    def _on_save(self):
        inv_num = self._vars['invoice_number'].get().strip()
        client  = self._vars['client_name'].get().strip()
        if not inv_num:
            messagebox.showwarning('Required', 'Invoice number is required.', parent=self)
            return
        if not client:
            messagebox.showwarning('Required', 'Client name is required.', parent=self)
            return
        # Check for duplicate invoice number
        existing = {r.get('invoice_number', '') for r in self.ds.read_invoices()}
        if inv_num in existing:
            if not messagebox.askyesno('Duplicate',
                    f'Invoice #{inv_num} already exists.\nOverwrite its record?', parent=self):
                return
        try:
            subtotal = float(self._subtotal_var.get().replace(',', '').replace('$', '') or 0)
            gst      = float(self._gst_var.get().replace(',', '').replace('$', '') or 0)
            total    = float(self._total_var.get().replace(',', '').replace('$', '') or 0)
        except ValueError:
            messagebox.showwarning('Invalid amount', 'Amounts must be numbers.', parent=self)
            return
        if total == 0 and subtotal > 0:
            total = subtotal + gst
        self.result = {
            'invoice_number':  inv_num,
            'invoice_date':    display_to_storage(self._vars['invoice_date'].get()),
            'due_date':        display_to_storage(self._vars['due_date'].get()),
            'client_name':     client,
            'client_address':  self._vars['client_address'].get().strip(),
            'notes':           self._notes_text.get('1.0', 'end').strip(),
            'subtotal':        f'{subtotal:.2f}',
            'gst':             f'{gst:.2f}',
            'total':           f'{total:.2f}',
            'paid':            'yes' if self._paid_var.get() else '',
            'paid_date':       display_to_storage(self._paid_date_var.get()),
            'payment_note':    self._pay_note_var.get().strip(),
            'invoice_status':  'paid' if self._paid_var.get() else 'unpaid',
            'pdf_path':        self._prefill.get('pdf_path', ''),
        }
        self.destroy()


class LoginDialog(tk.Toplevel):
    """Modal login dialog shown at startup."""

    def __init__(self, parent, ds):
        super().__init__(parent)
        self.title('Login — Invoice Generator')
        self.resizable(False, False)
        self.grab_set()
        self.ds = ds
        self.result = None
        self._build()
        self.protocol('WM_DELETE_WINDOW', self._on_cancel)
        self.wait_window(self)

    def _build(self):
        pad = {'padx': 12, 'pady': 6}
        ttk.Label(self, text='Username:').grid(row=0, column=0, sticky='e', **pad)
        self._user_var = tk.StringVar()
        ttk.Entry(self, textvariable=self._user_var, width=24).grid(row=0, column=1, sticky='w', **pad)
        ttk.Label(self, text='Password:').grid(row=1, column=0, sticky='e', **pad)
        self._pass_var = tk.StringVar()
        ttk.Entry(self, textvariable=self._pass_var, show='*', width=24).grid(row=1, column=1, sticky='w', **pad)
        btn = ttk.Frame(self)
        btn.grid(row=2, column=0, columnspan=2, pady=10)
        ttk.Button(btn, text='Login', command=self._on_login).pack(side='left', padx=4)
        ttk.Button(btn, text='Cancel', command=self._on_cancel).pack(side='left', padx=4)
        self._error_var = tk.StringVar()
        ttk.Label(self, textvariable=self._error_var, foreground='red').grid(row=3, column=0, columnspan=2, **pad)
        self.bind('<Return>', lambda e: self._on_login())

    def _on_login(self):
        user = self._user_var.get().strip()
        pw = self._pass_var.get()
        u = self.ds.authenticate_user(user, pw)
        if u:
            self.result = u
            self.destroy()
        else:
            self._error_var.set('Invalid username or password.')

    def _on_cancel(self):
        self.result = None
        self.destroy()


class UsersDialog(tk.Toplevel):
    """Manage users: add, edit, delete. No password requirements."""

    def __init__(self, parent, ds):
        super().__init__(parent)
        self.title('Manage Users')
        self.resizable(False, False)
        self.grab_set()
        self.ds = ds
        self._build()
        self._refresh()
        self.wait_window(self)

    def _build(self):
        cols = ('id', 'username', 'role', 'created_at')
        self._tree = ttk.Treeview(self, columns=cols, show='headings', selectmode='browse', height=10)
        self._tree.heading('id', text='ID')
        self._tree.heading('username', text='Username')
        self._tree.heading('role', text='Role')
        self._tree.heading('created_at', text='Created')
        self._tree.column('id', width=40, anchor='center')
        self._tree.column('username', width=140)
        self._tree.column('role', width=80, anchor='center')
        self._tree.column('created_at', width=130)
        self._tree.pack(fill='both', expand=True, padx=10, pady=10)
        frm = ttk.Frame(self)
        frm.pack(fill='x', padx=10, pady=(0, 10))
        ttk.Button(frm, text='Add', command=self._add).pack(side='left', padx=3)
        ttk.Button(frm, text='Edit Password…', command=self._edit).pack(side='left', padx=3)
        ttk.Button(frm, text='Delete', command=self._delete).pack(side='left', padx=3)
        ttk.Button(frm, text='Close', command=self.destroy).pack(side='right', padx=3)

    def _refresh(self):
        for i in self._tree.get_children():
            self._tree.delete(i)
        for u in self.ds.read_users():
            self._tree.insert('', 'end', values=(u.get('id', ''), u.get('username', ''),
                                                  u.get('role', ''), u.get('created_at', '')))

    def _selected_id(self):
        sel = self._tree.selection()
        if not sel:
            return None
        return self._tree.item(sel[0], 'values')[0]

    def _add(self):
        d = tk.Toplevel(self)
        d.title('Add User')
        d.resizable(False, False)
        d.grab_set()
        ttk.Label(d, text='Username:').grid(row=0, column=0, sticky='e', padx=8, pady=4)
        user_var = tk.StringVar()
        ttk.Entry(d, textvariable=user_var, width=22).grid(row=0, column=1, sticky='w', padx=8, pady=4)
        ttk.Label(d, text='Password:').grid(row=1, column=0, sticky='e', padx=8, pady=4)
        pass_var = tk.StringVar()
        ttk.Entry(d, textvariable=pass_var, show='*', width=22).grid(row=1, column=1, sticky='w', padx=8, pady=4)
        ttk.Label(d, text='Role:').grid(row=2, column=0, sticky='e', padx=8, pady=4)
        role_var = tk.StringVar(value='user')
        ttk.Combobox(d, textvariable=role_var, values=['user', 'admin'], state='readonly', width=18).grid(row=2, column=1, sticky='w', padx=8, pady=4)
        def save():
            u = user_var.get().strip()
            if not u:
                messagebox.showwarning('Required', 'Username is required.', parent=d)
                return
            if any(x.get('username', '').strip() == u for x in self.ds.read_users()):
                messagebox.showwarning('Duplicate', 'Username already exists.', parent=d)
                return
            self.ds.append_user({'username': u, 'password': pass_var.get(), 'role': role_var.get()})
            d.destroy()
            self._refresh()
        ttk.Button(d, text='Save', command=save).grid(row=3, column=1, sticky='e', padx=8, pady=8)

    def _edit(self):
        uid = self._selected_id()
        if not uid:
            return
        users = self.ds.read_users()
        u = next((x for x in users if x['id'] == uid), None)
        if not u:
            return
        d = tk.Toplevel(self)
        d.title(f'Edit User — {u.get("username", "")}')
        d.resizable(False, False)
        d.grab_set()
        ttk.Label(d, text='New password:').grid(row=0, column=0, sticky='e', padx=8, pady=6)
        pass_var = tk.StringVar(value=u.get('password', ''))
        ttk.Entry(d, textvariable=pass_var, show='*', width=22).grid(row=0, column=1, sticky='w', padx=8, pady=6)
        def save():
            self.ds.update_user(uid, {'password': pass_var.get()})
            d.destroy()
            self._refresh()
        ttk.Button(d, text='Save', command=save).grid(row=1, column=1, sticky='e', padx=8, pady=8)

    def _delete(self):
        uid = self._selected_id()
        if not uid:
            return
        u = next((x for x in self.ds.read_users() if x['id'] == uid), None)
        if not u:
            return
        if not messagebox.askyesno('Confirm', f"Delete user '{u.get('username', '')}'?", parent=self):
            return
        self.ds.delete_user(uid)
        self._refresh()


class PaymentDialog(tk.Toplevel):
    """Modal dialog for recording payment of an invoice."""

    def __init__(self, parent, invoice_number, client_name, total,
                 paid=False, paid_date='', payment_note=''):
        super().__init__(parent)
        self.title(f"Record Payment — Invoice #{invoice_number}")
        self.resizable(False, False)
        self.grab_set()
        self.result = None
        self._build(invoice_number, client_name, total, paid, paid_date, payment_note)
        self.wait_window(self)

    def _build(self, invoice_number, client_name, total, paid, paid_date, payment_note):
        pad = {'padx': 8, 'pady': 4}
        info = ttk.LabelFrame(self, text="Invoice")
        info.pack(fill='x', padx=12, pady=(12, 4))

        ttk.Label(info, text=f"Invoice #:").grid(row=0, column=0, sticky='e', **pad)
        ttk.Label(info, text=invoice_number, font=('TkDefaultFont', 10, 'bold')).grid(row=0, column=1, sticky='w', **pad)
        ttk.Label(info, text="Client:").grid(row=1, column=0, sticky='e', **pad)
        ttk.Label(info, text=client_name).grid(row=1, column=1, sticky='w', **pad)
        ttk.Label(info, text="Total:").grid(row=2, column=0, sticky='e', **pad)
        ttk.Label(info, text=total, font=('TkDefaultFont', 10, 'bold')).grid(row=2, column=1, sticky='w', **pad)

        details = ttk.LabelFrame(self, text="Payment Details")
        details.pack(fill='x', padx=12, pady=4)

        self._paid_var = tk.BooleanVar(value=paid)
        ttk.Checkbutton(details, text="Mark as paid", variable=self._paid_var).grid(
            row=0, column=0, columnspan=2, sticky='w', **pad)

        ttk.Label(details, text="Date paid:").grid(row=1, column=0, sticky='e', **pad)
        _paid_default = storage_to_display(paid_date) if paid_date else fmt_display(datetime.now().date())
        self._date_var = tk.StringVar(value=_paid_default)
        DateEntry(details, textvariable=self._date_var, width=12).grid(row=1, column=1, sticky='w', **pad)

        ttk.Label(details, text="Note:").grid(row=2, column=0, sticky='ne', **pad)
        self._note_text = tk.Text(details, width=35, height=4)
        self._note_text.grid(row=2, column=1, sticky='w', **pad)
        if payment_note:
            self._note_text.insert('1.0', payment_note)

        btn_frame = ttk.Frame(self)
        btn_frame.pack(fill='x', padx=12, pady=(4, 12))
        ttk.Button(btn_frame, text="Save", command=self._on_save).pack(side='right', padx=4)
        ttk.Button(btn_frame, text="Cancel", command=self.destroy).pack(side='right')

    def _on_save(self):
        self.result = {
            'paid':         self._paid_var.get(),
            'paid_date':    display_to_storage(self._date_var.get()),
            'payment_note': self._note_text.get('1.0', 'end').strip(),
        }
        self.destroy()



class InvoiceApp:
    """Main application class for invoice generation."""

    def __init__(self, root, ds=None, current_user='admin'):
        self.root = root
        self.root._app = self  # allow sibling tabs to find the app instance
        self.current_user = current_user
        self.root.title(f"Invoice Generator v{APP_VERSION}")
        self.root.minsize(1000, 750)
        self.root.geometry('1200x900')

        # Set up paths via DataStore
        self.base_path = Path(os.getcwd())
        self.ds = ds or DataStore(self.base_path)

        # Legacy path aliases (keep existing code working)
        self.settings_path      = self.ds.settings_path
        self.service_items_path = self.ds.service_items_path
        self.clients_path       = self.ds.clients_path
        self.invoices_dir       = self.ds.invoices_dir
        self.invoices_csv_path  = self.ds.invoices_csv_path

        # Ensure required directories/files exist
        self._ensure_environment()
        self.ds.ensure_files()
        self.ds.migrate_all()

        # Load settings, service items and clients
        self.settings = self._load_settings()
        self.service_items = self._load_service_items()
        self.clients = self._load_clients()

        # Data structures
        self.items = []  # list of dicts for each line item
        self._editing_index = None  # index of item being edited (None = adding new)

        # Build the user interface
        self._build_ui()

        # Populate invoice details
        self._populate_invoice_details()

        # Auto-backup scheduler
        from auto_backup import AutoBackupManager
        self._backup_mgr = AutoBackupManager(self.ds, lambda: self.settings)
        self._backup_mgr.start()
        self.root.protocol('WM_DELETE_WINDOW', self._on_close)

        # Periodic auto-refresh of all tabs
        self.root.after(5 * 60 * 1000, self._auto_refresh)

    # ------------------------------------------------------------------
    # Environment and data loading
    # ------------------------------------------------------------------
    def _ensure_environment(self):
        """Create necessary folders and files on first run."""
        # Create invoices directory
        self.invoices_dir.mkdir(exist_ok=True)
        # Create settings file if it doesn't exist
        if not self.settings_path.exists():
            with open(self.settings_path, 'w', encoding='utf-8') as f:
                json.dump(_DEFAULT_SETTINGS, f, indent=4)
        # Create default service items file if missing
        if not self.service_items_path.exists():
            with open(self.service_items_path, 'w', newline='', encoding='utf-8') as f:
                writer = csv.writer(f)
                writer.writerow(['description', 'unit_price', 'taxable'])
                writer.writerow(['First Aid Training', '300.00', 'yes'])
                writer.writerow(['Clinical Cover (per hour)', '100.00', 'yes'])
                writer.writerow(['Course Completion Certificate', '25.00', 'no'])
        # Create clients CSV if missing
        if not self.clients_path.exists():
            with open(self.clients_path, 'w', newline='', encoding='utf-8') as f:
                writer = csv.DictWriter(f, fieldnames=['name', 'contact_name', 'phone', 'email', 'address'])
                writer.writeheader()
        # Create invoices log CSV if missing
        if not self.invoices_csv_path.exists():
            with open(self.invoices_csv_path, 'w', newline='', encoding='utf-8') as f:
                writer = csv.writer(f)
                writer.writerow([
                    'invoice_number', 'invoice_date', 'due_date', 'client_name',
                    'client_address', 'notes', 'subtotal', 'gst', 'total',
                    'paid', 'paid_date', 'payment_note', 'invoice_status', 'pdf_path'
                ])
        else:
            self._migrate_invoices_csv()

    def _migrate_invoices_csv(self):
        """Add payment, status, and PDF path columns to existing invoices.csv if missing."""
        try:
            with open(self.invoices_csv_path, 'r', encoding='utf-8') as f:
                rows = list(csv.DictReader(f))
            if not rows:
                return
            # Full current field set
            fieldnames = [
                'invoice_number', 'invoice_date', 'due_date', 'client_name',
                'client_address', 'notes', 'subtotal', 'gst', 'total',
                'paid', 'paid_date', 'payment_note', 'invoice_status', 'pdf_path'
            ]
            missing = [f for f in fieldnames if f not in rows[0]]
            if not missing:
                return
            for row in rows:
                for f in fieldnames:
                    row.setdefault(f, '')
                # Derive canonical status if missing
                if 'invoice_status' in missing:
                    is_paid = row.get('paid', '').lower() in ('yes', 'true', '1')
                    row['invoice_status'] = 'paid' if is_paid else 'unpaid'
            with open(self.invoices_csv_path, 'w', newline='', encoding='utf-8') as f:
                writer = csv.DictWriter(f, fieldnames=fieldnames)
                writer.writeheader()
                writer.writerows(rows)
        except Exception:
            pass

    def _load_settings(self):
        """Load settings from JSON file, merging with defaults for missing keys."""
        try:
            with open(self.settings_path, 'r', encoding='utf-8') as f:
                loaded = json.load(f)
        except Exception:
            loaded = {}
        merged = dict(_DEFAULT_SETTINGS)
        merged.update(loaded)
        return merged

    def _save_settings(self):
        """Save settings back to JSON file."""
        with open(self.settings_path, 'w', encoding='utf-8') as f:
            json.dump(self.settings, f, indent=4)

    def _load_clients(self):
        """Read clients from CSV file into a list of dictionaries."""
        clients = []
        try:
            with open(self.clients_path, 'r', encoding='utf-8') as f:
                reader = csv.DictReader(f)
                for row in reader:
                    name = row.get('name', '').strip()
                    if name:
                        clients.append({
                            'name':         name,
                            'contact_name': row.get('contact_name', '').strip(),
                            'phone':        row.get('phone', '').strip(),
                            'email':        row.get('email', '').strip(),
                            'address':      row.get('address', '').strip(),
                        })
        except Exception:
            pass
        return clients

    def _load_service_items(self):
        """Read service items from CSV file into a list of dictionaries."""
        items = []
        try:
            with open(self.service_items_path, 'r', encoding='utf-8') as f:
                reader = csv.DictReader(f)
                for row in reader:
                    desc = row.get('description', '').strip()
                    price = float(row.get('unit_price', '0').strip() or '0')
                    taxable = row.get('taxable', '').strip().lower() in ('yes', 'true', '1')
                    if desc:
                        items.append({'description': desc, 'unit_price': price, 'taxable': taxable})
        except Exception:
            pass
        return items

    # ------------------------------------------------------------------
    # User Interface setup
    # ------------------------------------------------------------------
    def _build_ui(self):
        """Construct the user interface components."""
        # Menu bar
        menubar = tk.Menu(self.root)
        self.root.config(menu=menubar)

        edit_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Edit", menu=edit_menu)
        edit_menu.add_command(label="New Invoice",         command=self._new_invoice, accelerator='Ctrl+N')
        edit_menu.add_command(label="Save Invoice",        command=self._save_invoice, accelerator='Ctrl+S')
        edit_menu.add_separator()
        edit_menu.add_command(label="Add Ledger Entry…",   command=self._add_ledger_menu)
        edit_menu.add_separator()
        edit_menu.add_command(label="Settings…",           command=self._open_settings)
        edit_menu.add_command(label="Manage Clients…",     command=self._open_clients)
        edit_menu.add_command(label="Manage Service Catalogue…", command=self._open_catalogue)

        users_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Users", menu=users_menu)
        users_menu.add_command(label="Manage Users…", command=self._open_users)

        tools_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Tools", menu=tools_menu)
        tools_menu.add_command(label="Settings...", command=self._open_settings)
        tools_menu.add_command(label="Manage Clients...", command=self._open_clients)
        tools_menu.add_command(label="Manage Service Catalogue...", command=self._open_catalogue)
        tools_menu.add_separator()
        tools_menu.add_command(label="Open Invoices Folder", command=lambda: _open_file(self.invoices_dir))
        tools_menu.add_command(label="Open Data Folder", command=lambda: _open_file(self.ds.data_dir))
        tools_menu.add_separator()
        tools_menu.add_command(label="Reload All from Disk", command=self._reload_all,
                               accelerator='F5')
        tools_menu.add_separator()
        tools_menu.add_command(label='Backup Now',         command=self._backup_now)
        tools_menu.add_command(label="Export All Data...", command=self._export_data)
        tools_menu.add_command(label="Import Data...",    command=self._import_data)
        tools_menu.add_command(label="Migrate from V1.8 folder/zip...", command=self._migrate_from_v15)
        tools_menu.add_separator()
        tools_menu.add_command(label="Receipt Generator…", command=self._launch_receipt_generator)

        help_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Help", menu=help_menu)
        help_menu.add_command(label="Help Guide",
                              command=self._open_help_guide)
        help_menu.add_separator()
        help_menu.add_command(label="About Invoice Generator...",
                              command=lambda: AboutDialog(self.root))

        self.root.bind('<F5>', lambda e: self._reload_all())

        # Status bar (packed at bottom before notebook so it stays visible)
        self._build_status_bar()

        # Notebook
        self.notebook = ttk.Notebook(self.root)
        self.notebook.pack(fill='both', expand=True)

        create_tab  = ttk.Frame(self.notebook)
        history_tab = ttk.Frame(self.notebook)
        clients_tab = ttk.Frame(self.notebook)
        self.notebook.add(create_tab,  text='  Create Invoice  ')
        self.notebook.add(history_tab, text='  Invoice History  ')
        self.notebook.add(clients_tab, text='  Clients  ')
        self.notebook.bind('<<NotebookTabChanged>>', self._on_tab_changed)

        self._build_create_tab(create_tab)
        self._build_history_tab(history_tab)
        self._build_clients_tab(clients_tab)

        # New feature tabs — wired to DataStore
        self._ledger_tab   = LedgerTab(self.notebook, self.ds, lambda: self.settings.get('currency_symbol', '$'))
        self._reports_tab  = ReportsTab(self.notebook, self.ds,
                                        lambda: self.settings,
                                        lambda: self.settings.get('currency_symbol', '$'))

    def _build_status_bar(self):
        """Build the bottom status bar showing program location and live date/time."""
        bar = ttk.Frame(self.root, relief='sunken')
        bar.pack(side='bottom', fill='x')
        loc = str(Path(sys.executable).parent if getattr(sys, 'frozen', False)
                  else Path(__file__).parent)
        self._status_user_var = tk.StringVar(
            value=f"User: {getattr(self, 'current_user', 'admin')}")
        ttk.Label(bar, textvariable=self._status_user_var,
                  foreground=LABEL_MUTED).pack(side='left', padx=8, pady=2)
        ttk.Label(bar, text=f"Location: {loc}",
                  foreground=LABEL_MUTED).pack(side='left', padx=8, pady=2)
        self._status_clock_var = tk.StringVar()
        ttk.Label(bar, textvariable=self._status_clock_var,
                  foreground=LABEL_MUTED).pack(side='right', padx=8, pady=2)
        self._update_clock()

    def _update_clock(self):
        """Refresh the status bar clock every second."""
        try:
            self._status_clock_var.set(datetime.now().strftime('%a %d/%m/%Y  %H:%M:%S'))
            self.root.after(1000, self._update_clock)
        except Exception:
            pass

    def _auto_refresh(self):
        """Periodically reload data and refresh tabs so external changes appear."""
        try:
            interval_min = int(self.settings.get('auto_refresh_minutes', 5) or 5)
        except (ValueError, TypeError):
            interval_min = 5
        if interval_min > 0:
            try:
                self._history_refresh()
                self._clients_tab_refresh()
                for tab_obj in (
                    getattr(self, '_ledger_tab', None),
                    getattr(self, '_reports_tab', None),
                ):
                    if tab_obj and hasattr(tab_obj, 'refresh'):
                        try:
                            tab_obj.refresh()
                        except Exception:
                            pass
            except Exception:
                pass
            self.root.after(interval_min * 60 * 1000, self._auto_refresh)

    def _build_create_tab(self, parent):
        """Build the invoice creation tab with a scrollable canvas."""
        # Scrollable canvas wrapper
        canvas = tk.Canvas(parent, borderwidth=0, highlightthickness=0)
        vsb = ttk.Scrollbar(parent, orient='vertical', command=canvas.yview)
        canvas.configure(yscrollcommand=vsb.set)
        vsb.pack(side='right', fill='y')
        canvas.pack(side='left', fill='both', expand=True)

        inner = ttk.Frame(canvas)
        inner_id = canvas.create_window((0, 0), window=inner, anchor='nw')

        def _on_configure(event):
            canvas.configure(scrollregion=canvas.bbox('all'))
            canvas.itemconfig(inner_id, width=canvas.winfo_width())

        inner.bind('<Configure>', _on_configure)
        canvas.bind('<Configure>', lambda e: canvas.itemconfig(inner_id, width=canvas.winfo_width()))

        # Mouse-wheel scrolling
        def _on_mousewheel(event):
            canvas.yview_scroll(int(-1 * (event.delta / 120)), 'units')

        canvas.bind_all('<MouseWheel>', _on_mousewheel)

        self._build_create_inner(inner)

    def _build_create_inner(self, parent):
        """Build the actual create-invoice widgets inside the scrollable frame."""
        # Top frame for invoice metadata
        meta_frame = ttk.LabelFrame(parent, text="Invoice Details")
        meta_frame.grid(row=0, column=0, padx=10, pady=10, sticky='ew')
        parent.columnconfigure(0, weight=1)

        # Invoice number
        ttk.Label(meta_frame, text="Invoice #:").grid(row=0, column=0, sticky='e', padx=4, pady=3)
        self.invoice_number_var = tk.StringVar()
        self.invoice_number_entry = ttk.Entry(meta_frame, textvariable=self.invoice_number_var, state='readonly', width=15)
        self.invoice_number_entry.grid(row=0, column=1, sticky='w', padx=5)

        # Invoice date
        ttk.Label(meta_frame, text="Invoice date:").grid(row=1, column=0, sticky='e', padx=4, pady=3)
        self.invoice_date_var = tk.StringVar()
        DateEntry(meta_frame, textvariable=self.invoice_date_var, width=12).grid(row=1, column=1, sticky='w', padx=5)

        # Due date
        ttk.Label(meta_frame, text="Due date:").grid(row=2, column=0, sticky='e', padx=4, pady=3)
        self.due_date_var = tk.StringVar()
        DateEntry(meta_frame, textvariable=self.due_date_var, width=12).grid(row=2, column=1, sticky='w', padx=5)

        # Client picker from saved clients
        ttk.Label(meta_frame, text="Select client:").grid(row=3, column=0, sticky='e', padx=4, pady=3)
        self.client_pick_var = tk.StringVar()
        client_names = [''] + [c['name'] for c in self.clients]
        self.client_combo = ttk.Combobox(meta_frame, textvariable=self.client_pick_var,
                                         values=client_names, state='readonly', width=38)
        self.client_combo.grid(row=3, column=1, sticky='w', padx=5, columnspan=2)
        self.client_combo.bind('<<ComboboxSelected>>', self._client_selected)
        ttk.Button(meta_frame, text="Manage Clients",
                   command=self._open_clients).grid(row=3, column=3, sticky='w', padx=4)

        # Client name (manual override)
        ttk.Label(meta_frame, text="Client name:").grid(row=4, column=0, sticky='e', padx=4, pady=3)
        self.client_name_var = tk.StringVar()
        ttk.Entry(meta_frame, textvariable=self.client_name_var, width=40).grid(row=4, column=1, sticky='w', padx=5, columnspan=3)

        # Client address
        ttk.Label(meta_frame, text="Client address:").grid(row=5, column=0, sticky='ne', padx=4, pady=3)
        self.client_address_text = tk.Text(meta_frame, width=40, height=4)
        self.client_address_text.grid(row=5, column=1, sticky='w', padx=5, columnspan=3)

        # Notes
        ttk.Label(meta_frame, text="Notes:").grid(row=6, column=0, sticky='ne', padx=4, pady=3)
        self.notes_text = tk.Text(meta_frame, width=40, height=3)
        self.notes_text.grid(row=6, column=1, sticky='w', padx=5, columnspan=3)

        ttk.Separator(parent, orient='horizontal').grid(row=1, column=0, sticky='ew', padx=10, pady=5)

        # Item entry frame
        item_entry_frame = ttk.LabelFrame(parent, text="Add / Edit Line Item")
        item_entry_frame.grid(row=2, column=0, padx=10, pady=5, sticky='ew')

        # Service catalogue dropdown
        ttk.Label(item_entry_frame, text="Service:").grid(row=0, column=0, sticky='e', padx=4, pady=3)
        self.service_var = tk.StringVar()
        service_options = [''] + [item['description'] for item in self.service_items]
        self.service_combo = ttk.Combobox(item_entry_frame, textvariable=self.service_var, values=service_options, state='readonly', width=40)
        self.service_combo.grid(row=0, column=1, sticky='w', padx=5)
        self.service_combo.bind('<<ComboboxSelected>>', self._service_selected)

        # Description field
        ttk.Label(item_entry_frame, text="Description:").grid(row=1, column=0, sticky='e', padx=4, pady=3)
        self.desc_var = tk.StringVar()
        ttk.Entry(item_entry_frame, textvariable=self.desc_var, width=40).grid(row=1, column=1, sticky='w', padx=5)

        # Quantity
        ttk.Label(item_entry_frame, text="Qty:").grid(row=0, column=2, sticky='e', padx=4)
        self.qty_var = tk.StringVar()
        ttk.Entry(item_entry_frame, textvariable=self.qty_var, width=8).grid(row=0, column=3, sticky='w', padx=5)

        # Unit price
        ttk.Label(item_entry_frame, text="Unit price:").grid(row=1, column=2, sticky='e', padx=4)
        self.price_var = tk.StringVar()
        ttk.Entry(item_entry_frame, textvariable=self.price_var, width=8).grid(row=1, column=3, sticky='w', padx=5)

        # Taxable checkbox
        self.taxable_var = tk.BooleanVar()
        ttk.Checkbutton(item_entry_frame, text="Taxable", variable=self.taxable_var).grid(row=0, column=4, sticky='w', padx=5)

        # Add / Update item button
        self._add_btn_text = tk.StringVar(value="Add Item")
        ttk.Button(item_entry_frame, textvariable=self._add_btn_text, command=self._add_item).grid(row=1, column=4, sticky='w', padx=5)
        ttk.Button(item_entry_frame, text="Cancel Edit", command=self._cancel_edit).grid(row=1, column=5, sticky='w', padx=5)

        # Items table frame
        table_frame = ttk.LabelFrame(parent, text="Invoice Items  (double-click to edit)")
        table_frame.grid(row=3, column=0, padx=10, pady=5, sticky='nsew')
        parent.rowconfigure(3, weight=1)

        # Define columns for treeview
        sym = self.settings.get('currency_symbol', '$')
        columns = ('description', 'qty', 'unit_price', 'subtotal', 'taxable', 'gst', 'total')
        self.tree = ttk.Treeview(table_frame, columns=columns, show='headings', selectmode='browse', height=6)
        self.tree.grid(row=0, column=0, sticky='nsew')

        self.tree.heading('description', text='Description')
        self.tree.heading('qty', text='Qty')
        self.tree.heading('unit_price', text=f'Unit Price ({sym})')
        self.tree.heading('subtotal', text=f'Subtotal ({sym})')
        self.tree.heading('taxable', text='Taxable')
        self.tree.heading('gst', text=f'GST ({sym})')
        self.tree.heading('total', text=f'Total ({sym})')

        self.tree.column('description', width=200)
        self.tree.column('qty', width=40, anchor='center')
        self.tree.column('unit_price', width=90, anchor='e')
        self.tree.column('subtotal', width=90, anchor='e')
        self.tree.column('taxable', width=60, anchor='center')
        self.tree.column('gst', width=80, anchor='e')
        self.tree.column('total', width=90, anchor='e')

        scrollbar = ttk.Scrollbar(table_frame, orient='vertical', command=self.tree.yview)
        self.tree.configure(yscroll=scrollbar.set)
        scrollbar.grid(row=0, column=1, sticky='ns')
        table_frame.columnconfigure(0, weight=1)
        table_frame.rowconfigure(0, weight=1)

        # Double-click to edit
        self.tree.bind('<Double-1>', self._edit_selected)

        # Remove item button
        ttk.Button(table_frame, text="Remove Selected", command=self._remove_selected).grid(row=1, column=0, sticky='w', pady=5)

        # Totals frame
        gst_pct = int(self.settings.get('gst_rate', 0.10) * 100)
        totals_frame = ttk.Frame(parent)
        totals_frame.grid(row=4, column=0, padx=10, pady=5, sticky='e')

        ttk.Label(totals_frame, text="Subtotal:").grid(row=0, column=0, sticky='e', padx=4)
        self.subtotal_var = tk.StringVar(value=f"{sym}0.00")
        ttk.Label(totals_frame, textvariable=self.subtotal_var, width=14, anchor='e', font=('TkDefaultFont', 10)).grid(row=0, column=1, sticky='e')

        ttk.Label(totals_frame, text=f"GST ({gst_pct}%):").grid(row=1, column=0, sticky='e', padx=4)
        self.gst_var = tk.StringVar(value=f"{sym}0.00")
        ttk.Label(totals_frame, textvariable=self.gst_var, width=14, anchor='e', font=('TkDefaultFont', 10)).grid(row=1, column=1, sticky='e')

        ttk.Label(totals_frame, text="Total:", font=('TkDefaultFont', 11, 'bold')).grid(row=2, column=0, sticky='e', padx=4)
        self.total_var = tk.StringVar(value=f"{sym}0.00")
        ttk.Label(totals_frame, textvariable=self.total_var, width=14, anchor='e', font=('TkDefaultFont', 11, 'bold')).grid(row=2, column=1, sticky='e')

        # Action buttons
        button_frame = ttk.Frame(parent)
        button_frame.grid(row=5, column=0, padx=10, pady=10, sticky='e')

        ttk.Button(button_frame, text="Save Invoice", command=self._save_invoice).grid(row=0, column=0, padx=5)
        ttk.Button(button_frame, text="Generate Draft", command=self._generate_draft).grid(row=0, column=1, padx=5)
        ttk.Button(button_frame, text="Clear", command=self._clear_form).grid(row=0, column=2, padx=5)
        ttk.Button(button_frame, text="Reveal Invoice Folder", command=self._reveal_invoice_folder).grid(row=0, column=3, padx=5)

    def _build_clients_tab(self, parent):
        """Build the Clients overview tab."""
        # Summary treeview
        cols = ('name', 'contact_name', 'phone', 'email', 'invoices', 'total_billed')
        self.clients_tree = ttk.Treeview(parent, columns=cols, show='headings',
                                         selectmode='browse', height=16)
        self.clients_tree.heading('name',          text='Client Name')
        self.clients_tree.heading('contact_name',  text='Contact')
        self.clients_tree.heading('phone',         text='Phone')
        self.clients_tree.heading('email',         text='Email')
        self.clients_tree.heading('invoices',      text='Invoices')
        self.clients_tree.heading('total_billed',  text='Total Billed')
        self.clients_tree.column('name',         width=160)
        self.clients_tree.column('contact_name', width=120)
        self.clients_tree.column('phone',        width=100)
        self.clients_tree.column('email',        width=160)
        self.clients_tree.column('invoices',     width=65,  anchor='center')
        self.clients_tree.column('total_billed', width=100, anchor='e')
        sb = ttk.Scrollbar(parent, orient='vertical', command=self.clients_tree.yview)
        self.clients_tree.configure(yscroll=sb.set)
        self.clients_tree.pack(side='left', fill='both', expand=True, padx=(10, 0), pady=10)
        sb.pack(side='left', fill='y', pady=10)
        bind_treeview_clipboard(self.clients_tree, self.root)

        btn_frame = ttk.Frame(parent)
        btn_frame.pack(side='right', fill='y', padx=10, pady=10)
        ttk.Button(btn_frame, text="Manage Clients",
                   command=self._open_clients).pack(pady=3, fill='x')
        ttk.Button(btn_frame, text="Refresh",
                   command=self._clients_tab_refresh).pack(pady=3, fill='x')

    def _build_history_tab(self, parent):
        """Build the invoice history tab."""
        columns = ('number', 'date', 'due_date', 'client', 'total', 'status')
        self.hist_tree = ttk.Treeview(parent, columns=columns, show='headings', selectmode='browse', height=15)
        self.hist_tree.heading('number',   text='Invoice #')
        self.hist_tree.heading('date',     text='Date')
        self.hist_tree.heading('due_date', text='Due Date')
        self.hist_tree.heading('client',   text='Client')
        self.hist_tree.heading('total',    text='Total')
        self.hist_tree.heading('status',   text='Status')
        self.hist_tree.column('number',   width=75,  anchor='center')
        self.hist_tree.column('date',     width=95,  anchor='center')
        self.hist_tree.column('due_date', width=95,  anchor='center')
        self.hist_tree.column('client',   width=185)
        self.hist_tree.column('total',    width=90,  anchor='e')
        self.hist_tree.column('status',   width=80,  anchor='center')
        self.hist_tree.tag_configure('paid',      background='#d4edda', foreground='#155724')
        self.hist_tree.tag_configure('unpaid',    background='#fff3cd', foreground='#856404')
        self.hist_tree.tag_configure('cancelled', background='#f8d7da', foreground='#721c24')
        self.hist_tree.tag_configure('void',      background='#e2e3e5', foreground='#6c757d')
        sb = ttk.Scrollbar(parent, orient='vertical', command=self.hist_tree.yview)
        self.hist_tree.configure(yscroll=sb.set)
        self.hist_tree.pack(side='left', fill='both', expand=True, padx=(10, 0), pady=10)
        sb.pack(side='left', fill='y', pady=10)
        bind_treeview_clipboard(self.hist_tree, self.root)

        btn_frame = ttk.Frame(parent)
        btn_frame.pack(side='right', fill='y', padx=10, pady=10)
        ttk.Button(btn_frame, text="Open PDF",            command=self._history_open_pdf).pack(pady=3, fill='x')
        ttk.Button(btn_frame, text="Edit Invoice…",        command=self._edit_invoice).pack(pady=3, fill='x')
        ttk.Button(btn_frame, text="Record Payment",       command=self._record_payment).pack(pady=3, fill='x')
        ttk.Button(btn_frame, text="Cancel Invoice",       command=self._cancel_invoice).pack(pady=3, fill='x')
        ttk.Button(btn_frame, text="Reissue Invoice",      command=self._reissue_invoice).pack(pady=3, fill='x')
        ttk.Button(btn_frame, text="Link PDF…",            command=self._link_invoice_pdf).pack(pady=3, fill='x')
        ttk.Button(btn_frame, text="Reveal Invoice Folder", command=self._reveal_invoice_folder).pack(pady=3, fill='x')
        ttk.Button(btn_frame, text="Record Missing…",      command=self._record_missing_invoice).pack(pady=3, fill='x')
        ttk.Button(btn_frame, text="Copy to Clipboard",    command=lambda: self.hist_tree.event_generate('<Control-Shift-C>')).pack(pady=3, fill='x')
        ttk.Button(btn_frame, text="Export to CSV…",       command=self._export_history_csv).pack(pady=3, fill='x')
        ttk.Button(btn_frame, text="Refresh",              command=self._history_refresh).pack(pady=3, fill='x')
        self.hist_tree.bind('<Double-1>', lambda e: self._edit_invoice())

    # ------------------------------------------------------------------
    # Invoice initialisation
    # ------------------------------------------------------------------
    def _populate_invoice_details(self):
        """Fill invoice number and default dates."""
        next_inv = self.settings.get('next_invoice_number', 1)
        self.invoice_number_var.set(f"{next_inv:04d}")
        today = datetime.now().date()
        self.invoice_date_var.set(fmt_display(today))
        days = self.settings.get('payment_terms_days', 30)
        due = today + timedelta(days=days)
        self.due_date_var.set(fmt_display(due))

    def _prompt_past_courses(self):
        """Course-related feature removed; kept as no-op for compatibility."""
        pass

    # ------------------------------------------------------------------
    # Item management
    # ------------------------------------------------------------------
    def _service_selected(self, event=None):
        """Populate description, price and tax when a service is selected."""
        selected = self.service_var.get()
        if not selected:
            return
        for item in self.service_items:
            if item['description'] == selected:
                self.desc_var.set(item['description'])
                self.price_var.set(f"{item['unit_price']:.2f}")
                self.taxable_var.set(item['taxable'])
                if not self.qty_var.get():
                    self.qty_var.set('1')
                break

    def _edit_selected(self, event=None):
        """Load a line item into the entry fields for editing (double-click)."""
        selected = self.tree.selection()
        if not selected:
            return
        index = self.tree.index(selected[0])
        item = self.items[index]
        self._editing_index = index
        self.desc_var.set(item['description'])
        self.qty_var.set(f"{item['qty']:g}")
        self.price_var.set(f"{item['unit_price']:.2f}")
        self.taxable_var.set(item['taxable'])
        self.service_var.set('')
        self._add_btn_text.set("Update Item")

    def _cancel_edit(self):
        """Cancel an in-progress edit and reset entry fields."""
        self._editing_index = None
        self._add_btn_text.set("Add Item")
        self.service_var.set('')
        self.desc_var.set('')
        self.qty_var.set('')
        self.price_var.set('')
        self.taxable_var.set(False)

    def _add_item(self):
        """Validate inputs and add or update a line item in the table."""
        desc = self.desc_var.get().strip()
        qty_str = self.qty_var.get().strip()
        price_str = self.price_var.get().strip()
        taxable = self.taxable_var.get()
        if not desc:
            messagebox.showwarning("Missing description", "Please enter an item description.")
            return
        try:
            qty = float(qty_str)
        except ValueError:
            messagebox.showwarning("Invalid quantity", "Quantity must be a number.")
            return
        try:
            unit_price = float(price_str)
        except ValueError:
            messagebox.showwarning("Invalid unit price", "Unit price must be a number.")
            return
        if qty <= 0 or unit_price < 0:
            messagebox.showwarning("Invalid values", "Quantity must be positive and price non‑negative.")
            return

        gst_rate = self.settings.get('gst_rate', 0.10)
        subtotal = qty * unit_price
        gst = subtotal * gst_rate if taxable else 0.0
        total = subtotal + gst
        sym = self.settings.get('currency_symbol', '$')

        item = {
            'description': desc,
            'qty': qty,
            'unit_price': unit_price,
            'subtotal': subtotal,
            'taxable': taxable,
            'gst': gst,
            'total': total
        }

        if self._editing_index is not None:
            # Update existing row
            self.items[self._editing_index] = item
            iid = self.tree.get_children()[self._editing_index]
            self.tree.item(iid, values=(
                desc, f"{qty:g}",
                f"{sym}{unit_price:.2f}", f"{sym}{subtotal:.2f}",
                'Yes' if taxable else 'No',
                f"{sym}{gst:.2f}", f"{sym}{total:.2f}"
            ))
            self._editing_index = None
            self._add_btn_text.set("Add Item")
        else:
            self.items.append(item)
            self.tree.insert('', 'end', values=(
                desc, f"{qty:g}",
                f"{sym}{unit_price:.2f}", f"{sym}{subtotal:.2f}",
                'Yes' if taxable else 'No',
                f"{sym}{gst:.2f}", f"{sym}{total:.2f}"
            ))

        self._update_totals()
        self.service_var.set('')
        self.desc_var.set('')
        self.qty_var.set('')
        self.price_var.set('')
        self.taxable_var.set(False)

    def _remove_selected(self):
        """Remove the selected line item from the table and internal list."""
        selected = self.tree.selection()
        if not selected:
            return
        index = self.tree.index(selected[0])
        self.tree.delete(selected[0])
        try:
            self.items.pop(index)
        except IndexError:
            pass
        if self._editing_index == index:
            self._cancel_edit()
        self._update_totals()

    def _update_totals(self):
        """Recalculate and display subtotal, GST and total."""
        sym = self.settings.get('currency_symbol', '$')
        subtotal = sum(item['subtotal'] for item in self.items)
        gst = sum(item['gst'] for item in self.items)
        total = sum(item['total'] for item in self.items)
        self.subtotal_var.set(f"{sym}{subtotal:.2f}")
        self.gst_var.set(f"{sym}{gst:.2f}")
        self.total_var.set(f"{sym}{total:.2f}")

    # ------------------------------------------------------------------
    # Invoice saving and PDF generation
    # ------------------------------------------------------------------
    def _save_invoice(self):
        """Validate the invoice and save it to records and PDF."""
        if not self.items:
            messagebox.showwarning("No items", "Add at least one line item before saving.")
            return
        client_name = self.client_name_var.get().strip()
        if not client_name:
            messagebox.showwarning("Missing client", "Please enter the client's name.")
            return
        invoice_number = self.invoice_number_var.get().strip()
        invoice_date = display_to_storage(self.invoice_date_var.get())
        due_date = display_to_storage(self.due_date_var.get())
        client_address = self.client_address_text.get('1.0', 'end').strip()
        notes = self.notes_text.get('1.0', 'end').strip()
        sym = self.settings.get('currency_symbol', '$')
        subtotal = sum(i['subtotal'] for i in self.items)
        gst = sum(i['gst'] for i in self.items)
        total = sum(i['total'] for i in self.items)

        # Determine PDF save path
        pdf_filename = f"invoice_{invoice_number}.pdf"
        save_mode = self.settings.get('pdf_save_mode', 'auto')
        save_dir_setting = self.settings.get('pdf_save_dir', '').strip()
        default_dir = Path(save_dir_setting) if save_dir_setting else self.invoices_dir

        if save_mode == 'prompt':
            pdf_path = filedialog.asksaveasfilename(
                title="Save Invoice PDF",
                initialdir=str(default_dir),
                initialfile=pdf_filename,
                defaultextension='.pdf',
                filetypes=[('PDF files', '*.pdf'), ('All files', '*.*')]
            )
            if not pdf_path:  # user cancelled
                return
            pdf_path = Path(pdf_path)
        else:
            default_dir.mkdir(parents=True, exist_ok=True)
            pdf_path = default_dir / pdf_filename

        # Append to invoices.csv via DataStore
        self.ds.append_invoice({
            'invoice_number': invoice_number, 'invoice_date': invoice_date,
            'due_date': due_date, 'client_name': client_name,
            'client_address': client_address, 'notes': notes,
            'subtotal': f'{subtotal:.2f}', 'gst': f'{gst:.2f}', 'total': f'{total:.2f}',
            'paid': '', 'paid_date': '', 'payment_note': '',
            'invoice_status': 'unpaid', 'pdf_path': str(pdf_path) if save_mode == 'auto' else '',
        })

        success = self._create_pdf(
            pdf_path,
            invoice_number=invoice_number,
            invoice_date=invoice_date,
            due_date=due_date,
            client_name=client_name,
            client_address=client_address,
            notes=notes,
            items=self.items,
            subtotal=subtotal,
            gst=gst,
            total=total
        )

        # Update settings for next invoice
        self.settings['next_invoice_number'] = int(invoice_number) + 1
        self._save_settings()

        # Auto-post a matching ledger 'out' (receivable) entry
        try:
            self.ds.append_ledger({
                'date':        invoice_date or datetime.now().strftime('%Y-%m-%d'),
                'type':        'in',
                'category':    'Invoice Payment',
                'description': f'Invoice {invoice_number} — {client_name}',
                'amount':      f'{total:.2f}',
                'reference':   invoice_number,
                'notes':       'Auto-posted on invoice save (unpaid)',
            })
        except Exception:
            pass

        messagebox.showinfo(
            "Invoice Saved",
            f"Invoice {invoice_number} saved.\nPDF: {pdf_path}"
        )

        # Open the PDF automatically
        if success and pdf_path.exists():
            _open_file(pdf_path)

        # Reset form for next invoice
        self._clear_form(reset_invoice_number=True)

    def _create_pdf(self, filename, **kw):
        """Generate a PDF invoice using reportlab.

        Returns True on success, False on failure.
        """
        s = self.settings
        sym = s.get('currency_symbol', '$')
        gst_rate = s.get('gst_rate', 0.10)
        gst_pct = int(round(gst_rate * 100))

        doc = SimpleDocTemplate(
            str(filename), pagesize=A4,
            rightMargin=20 * mm, leftMargin=20 * mm,
            topMargin=20 * mm, bottomMargin=20 * mm
        )
        elements = []
        styles = getSampleStyleSheet()
        right_style = ParagraphStyle('right', parent=styles['Normal'], alignment=TA_RIGHT)

        # ---- Header: business left, invoice info right ----
        biz_lines = []
        if s.get('business_name'):
            biz_lines.append(f"<b>{s['business_name']}</b>")
        if s.get('business_abn'):
            biz_lines.append(f"ABN: {s['business_abn']}")
        if s.get('business_address'):
            for ln in s['business_address'].split('\n'):
                if ln.strip():
                    biz_lines.append(ln.strip())
        if s.get('business_phone'):
            biz_lines.append(f"Ph: {s['business_phone']}")
        if s.get('business_email'):
            biz_lines.append(s['business_email'])

        inv_lines = [
            f"<b>INVOICE</b>",
            f"<b>Invoice #:</b> {kw['invoice_number']}",
            f"<b>Date:</b> {kw['invoice_date']}",
            f"<b>Due:</b> {kw['due_date']}",
        ]

        header_data = [[
            Paragraph('<br/>'.join(biz_lines), styles['Normal']),
            Paragraph('<br/>'.join(inv_lines), right_style)
        ]]
        header_table = Table(header_data, colWidths=[90 * mm, 80 * mm])
        header_table.setStyle(TableStyle([
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ]))
        elements.append(header_table)
        elements.append(HRFlowable(width='100%', thickness=1, color=colors.black))
        elements.append(Spacer(1, 10))

        # ---- Billed to ----
        elements.append(Paragraph("<b>Billed To:</b>", styles['Normal']))
        elements.append(Paragraph(kw['client_name'], styles['Normal']))
        for line in (kw['client_address'] or '').split('\n'):
            if line.strip():
                elements.append(Paragraph(line.strip(), styles['Normal']))
        elements.append(Spacer(1, 12))

        # ---- Notes ----
        if kw.get('notes'):
            elements.append(Paragraph(f"<b>Notes:</b> {kw['notes']}", styles['Normal']))
            elements.append(Spacer(1, 12))

        # ---- Items table ----
        # Styles for wrapped cell content
        cell_style = ParagraphStyle('cell', parent=styles['Normal'], fontSize=9, leading=11)
        cell_right = ParagraphStyle('cell_r', parent=cell_style, alignment=TA_RIGHT)
        cell_center = ParagraphStyle('cell_c', parent=cell_style, alignment=TA_CENTER)
        hdr_style = ParagraphStyle('hdr', parent=styles['Normal'], fontSize=9,
                                    leading=11, textColor=colors.white, fontName='Helvetica-Bold')
        hdr_right = ParagraphStyle('hdr_r', parent=hdr_style, alignment=TA_RIGHT)
        hdr_center = ParagraphStyle('hdr_c', parent=hdr_style, alignment=TA_CENTER)

        data = [[
            Paragraph('Description', hdr_style),
            Paragraph('Qty', hdr_center),
            Paragraph(f'Unit Price<br/>({sym})', hdr_right),
            Paragraph(f'Subtotal<br/>({sym})', hdr_right),
            Paragraph('Taxable', hdr_center),
            Paragraph(f'GST<br/>({sym})', hdr_right),
            Paragraph(f'Total<br/>({sym})', hdr_right),
        ]]
        for item in kw['items']:
            data.append([
                Paragraph(item['description'], cell_style),
                Paragraph(f"{item['qty']:g}", cell_center),
                Paragraph(f"{sym}{item['unit_price']:.2f}", cell_right),
                Paragraph(f"{sym}{item['subtotal']:.2f}", cell_right),
                Paragraph('Yes' if item['taxable'] else 'No', cell_center),
                Paragraph(f"{sym}{item['gst']:.2f}", cell_right),
                Paragraph(f"{sym}{item['total']:.2f}", cell_right),
            ])
        data.append([
            Paragraph('', cell_style), Paragraph('', cell_style), Paragraph('', cell_right),
            Paragraph(f"<b>{sym}{kw['subtotal']:.2f}</b>", cell_right),
            Paragraph('', cell_style),
            Paragraph(f"<b>{sym}{kw['gst']:.2f}</b>", cell_right),
            Paragraph(f"<b>{sym}{kw['total']:.2f}</b>", cell_right),
        ])

        table = Table(data, colWidths=[68 * mm, 13 * mm, 24 * mm, 24 * mm, 18 * mm, 22 * mm, 24 * mm])
        t_style = TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4472C4')),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
            ('TOPPADDING', (0, 0), (-1, -1), 4),
            ('LEFTPADDING', (0, 0), (-1, -1), 4),
            ('RIGHTPADDING', (0, 0), (-1, -1), 4),
            ('GRID', (0, 0), (-1, -2), 0.5, colors.grey),
            ('ROWBACKGROUNDS', (0, 1), (-1, -2), [colors.white, colors.HexColor('#EEF2FF')]),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('LINEABOVE', (0, -1), (-1, -1), 1, colors.black),
            ('BACKGROUND', (0, -1), (-1, -1), colors.HexColor('#D9E1F2')),
        ])
        table.setStyle(t_style)
        elements.append(table)
        elements.append(Spacer(1, 14))

        # ---- Totals summary (right-aligned) ----
        totals_data = [
            ['Subtotal:', f"{sym}{kw['subtotal']:.2f}"],
            [f'GST ({gst_pct}%):', f"{sym}{kw['gst']:.2f}"],
            ['Total Due:', f"{sym}{kw['total']:.2f}"],
        ]
        totals_table = Table(totals_data, colWidths=[40 * mm, 30 * mm],
                             hAlign='RIGHT')
        totals_table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (0, -1), 'RIGHT'),
            ('ALIGN', (1, 0), (1, -1), 'RIGHT'),
            ('FONTNAME', (0, -1), (-1, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, -1), (-1, -1), 11),
            ('LINEABOVE', (0, -1), (-1, -1), 1, colors.black),
            ('TOPPADDING', (0, -1), (-1, -1), 4),
        ]))
        elements.append(totals_table)

        # ---- Payment details ----
        any_bank = any(s.get(k) for k in ('bank_name', 'bank_bsb', 'bank_account', 'bank_account_name'))
        if any_bank:
            elements.append(Spacer(1, 16))
            elements.append(HRFlowable(width='100%', thickness=0.5, color=colors.grey))
            elements.append(Spacer(1, 6))
            elements.append(Paragraph("<b>Payment Details</b>", styles['Normal']))
            if s.get('bank_account_name'):
                elements.append(Paragraph(f"Account Name: {s['bank_account_name']}", styles['Normal']))
            if s.get('bank_name'):
                elements.append(Paragraph(f"Bank: {s['bank_name']}", styles['Normal']))
            if s.get('bank_bsb'):
                elements.append(Paragraph(f"BSB: {s['bank_bsb']}", styles['Normal']))
            if s.get('bank_account'):
                elements.append(Paragraph(f"Account #: {s['bank_account']}", styles['Normal']))
            elements.append(Paragraph(
                f"Please use Invoice #{kw['invoice_number']} as payment reference.",
                styles['Normal']
            ))

        # ---- GST not-registered note ----
        if s.get('show_gst_not_registered'):
            elements.append(Spacer(1, 10))
            elements.append(Paragraph(
                "<i>This business is not registered for GST. No GST has been charged.</i>",
                styles['Normal']
            ))

        # ---- Thank-you note ----
        thank_you = s.get('thank_you_note', '').strip()
        if thank_you:
            elements.append(Spacer(1, 14))
            elements.append(HRFlowable(width='100%', thickness=0.5, color=colors.grey))
            elements.append(Spacer(1, 6))
            centre_style = ParagraphStyle('centre', parent=styles['Normal'], alignment=TA_CENTER,
                                          fontSize=10, textColor=colors.HexColor('#4472C4'))
            elements.append(Paragraph(thank_you, centre_style))

        try:
            doc.build(elements)
            return True
        except Exception as e:
            messagebox.showerror("PDF Error", f"Failed to create PDF:\n{e}")
            return False

    def _generate_draft(self):
        """Generate a .docx draft of the current invoice without recording it."""
        if not self.items:
            messagebox.showwarning("No items", "Add at least one line item before generating a draft.")
            return
        client_name = self.client_name_var.get().strip()
        if not client_name:
            messagebox.showwarning("Missing client", "Please enter the client's name.")
            return
        invoice_number = self.invoice_number_var.get().strip()
        invoice_date = display_to_storage(self.invoice_date_var.get())
        due_date = display_to_storage(self.due_date_var.get())
        client_address = self.client_address_text.get('1.0', 'end').strip()
        notes = self.notes_text.get('1.0', 'end').strip()
        sym = self.settings.get('currency_symbol', '$')
        subtotal = sum(i['subtotal'] for i in self.items)
        gst = sum(i['gst'] for i in self.items)
        total = sum(i['total'] for i in self.items)

        # Choose save location
        save_dir_setting = self.settings.get('pdf_save_dir', '').strip()
        default_dir = Path(save_dir_setting) if save_dir_setting else self.invoices_dir
        default_dir.mkdir(parents=True, exist_ok=True)
        docx_filename = f"invoice_{invoice_number}_draft.docx"

        docx_path = filedialog.asksaveasfilename(
            title="Save Draft (Word Document)",
            initialdir=str(default_dir),
            initialfile=docx_filename,
            defaultextension='.docx',
            filetypes=[('Word Document', '*.docx'), ('All files', '*.*')]
        )
        if not docx_path:
            return
        docx_path = Path(docx_path)

        try:
            self._create_docx(
                docx_path,
                invoice_number=invoice_number,
                invoice_date=invoice_date,
                due_date=due_date,
                client_name=client_name,
                client_address=client_address,
                notes=notes,
                items=self.items,
                subtotal=subtotal,
                gst=gst,
                total=total,
            )
            messagebox.showinfo("Draft Saved", f"Draft saved to:\n{docx_path}")
            _open_file(docx_path)
        except Exception as e:
            messagebox.showerror("Draft Error", f"Failed to create draft:\n{e}")

    def _create_docx(self, filename, **kw):
        """Generate a .docx draft invoice using python-docx."""
        from docx import Document
        from docx.shared import Pt, Inches, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT

        s = self.settings
        sym = s.get('currency_symbol', '$')
        gst_rate = s.get('gst_rate', 0.10)
        gst_pct = int(round(gst_rate * 100))

        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(10)

        # Header — business info
        if s.get('business_name'):
            p = doc.add_paragraph()
            run = p.add_run(s['business_name'])
            run.bold = True
            run.font.size = Pt(14)
        if s.get('business_abn'):
            doc.add_paragraph(f"ABN: {s['business_abn']}")
        if s.get('business_address'):
            for ln in s['business_address'].split('\n'):
                if ln.strip():
                    doc.add_paragraph(ln.strip())
        if s.get('business_phone'):
            doc.add_paragraph(f"Ph: {s['business_phone']}")
        if s.get('business_email'):
            doc.add_paragraph(s['business_email'])

        doc.add_paragraph()  # spacer

        # Invoice info
        p = doc.add_paragraph()
        run = p.add_run('INVOICE')
        run.bold = True
        run.font.size = Pt(16)
        doc.add_paragraph(f"Invoice #: {kw['invoice_number']}")
        doc.add_paragraph(f"Date: {kw['invoice_date']}")
        doc.add_paragraph(f"Due: {kw['due_date']}")

        doc.add_paragraph()  # spacer

        # Billed to
        p = doc.add_paragraph()
        p.add_run('Billed To:').bold = True
        doc.add_paragraph(kw['client_name'])
        if kw['client_address']:
            for line in kw['client_address'].split('\n'):
                if line.strip():
                    doc.add_paragraph(line.strip())

        # Notes
        if kw.get('notes'):
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.add_run('Notes: ').bold = True
            p.add_run(kw['notes'])

        doc.add_paragraph()  # spacer

        # Items table
        headers = ['Description', 'Qty', f'Unit Price ({sym})',
                   f'Subtotal ({sym})', 'Taxable', f'GST ({sym})', f'Total ({sym})']
        table = doc.add_table(rows=1, cols=7)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Set column widths (description gets the most space)
        col_widths = [Cm(7.5), Cm(1.5), Cm(2.5), Cm(2.5), Cm(1.8), Cm(2.2), Cm(2.5)]
        for i, width in enumerate(col_widths):
            table.columns[i].width = width

        # Header row
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(9)

        # Item rows
        for item in kw['items']:
            row = table.add_row().cells
            row[0].text = item['description']
            row[1].text = f"{item['qty']:g}"
            row[2].text = f"{sym}{item['unit_price']:.2f}"
            row[3].text = f"{sym}{item['subtotal']:.2f}"
            row[4].text = 'Yes' if item['taxable'] else 'No'
            row[5].text = f"{sym}{item['gst']:.2f}"
            row[6].text = f"{sym}{item['total']:.2f}"

        # Totals row
        row = table.add_row().cells
        row[0].text = ''
        row[2].text = ''
        row[3].text = f"{sym}{kw['subtotal']:.2f}"
        row[5].text = f"{sym}{kw['gst']:.2f}"
        row[6].text = f"{sym}{kw['total']:.2f}"
        for cell in row:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        doc.add_paragraph()  # spacer

        # Totals summary
        doc.add_paragraph(f"Subtotal: {sym}{kw['subtotal']:.2f}")
        doc.add_paragraph(f"GST ({gst_pct}%): {sym}{kw['gst']:.2f}")
        p = doc.add_paragraph()
        run = p.add_run(f"Total Due: {sym}{kw['total']:.2f}")
        run.bold = True
        run.font.size = Pt(12)

        # Payment details
        any_bank = any(s.get(k) for k in ('bank_name', 'bank_bsb', 'bank_account', 'bank_account_name'))
        if any_bank:
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.add_run('Payment Details').bold = True
            if s.get('bank_account_name'):
                doc.add_paragraph(f"Account Name: {s['bank_account_name']}")
            if s.get('bank_name'):
                doc.add_paragraph(f"Bank: {s['bank_name']}")
            if s.get('bank_bsb'):
                doc.add_paragraph(f"BSB: {s['bank_bsb']}")
            if s.get('bank_account'):
                doc.add_paragraph(f"Account #: {s['bank_account']}")
            doc.add_paragraph(
                f"Please use Invoice #{kw['invoice_number']} as payment reference.")

        # GST not-registered note
        if s.get('show_gst_not_registered'):
            doc.add_paragraph()
            p = doc.add_paragraph()
            run = p.add_run(
                'This business is not registered for GST. No GST has been charged.')
            run.italic = True

        # Thank-you note
        thank_you = s.get('thank_you_note', '').strip()
        if thank_you:
            doc.add_paragraph()
            p = doc.add_paragraph(thank_you)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.save(str(filename))

    # ------------------------------------------------------------------
    # Form management
    # ------------------------------------------------------------------
    def _clear_form(self, reset_invoice_number=False):
        """Clear the current invoice form for a new entry."""
        for item in self.tree.get_children():
            self.tree.delete(item)
        self.items.clear()
        self._editing_index = None
        self._add_btn_text.set("Add Item")
        self.client_pick_var.set('')
        self.client_name_var.set('')
        self.client_address_text.delete('1.0', 'end')
        self.notes_text.delete('1.0', 'end')
        self.service_var.set('')
        self.desc_var.set('')
        self.qty_var.set('')
        self.price_var.set('')
        self.taxable_var.set(False)
        sym = self.settings.get('currency_symbol', '$')
        self.subtotal_var.set(f"{sym}0.00")
        self.gst_var.set(f"{sym}0.00")
        self.total_var.set(f"{sym}0.00")
        self._populate_invoice_details()
        if not reset_invoice_number:
            current_inv = self.invoice_number_var.get().strip()
            if current_inv.isdigit():
                self.invoice_number_var.set(current_inv)

    # ------------------------------------------------------------------
    # History tab
    # ------------------------------------------------------------------
    def _on_tab_changed(self, event=None):
        """Refresh data-driven tabs when the user switches to them."""
        idx = self.notebook.index('current')
        if idx == 1:
            self._history_refresh()
        elif idx == 2:
            self._clients_tab_refresh()
        elif idx == 3:
            self._ledger_tab.refresh()
        elif idx == 4:
            self._reports_tab.refresh()

    def _history_refresh(self):
        """Load invoice history from CSV into the history treeview."""
        for row in self.hist_tree.get_children():
            self.hist_tree.delete(row)
        sym = self.settings.get('currency_symbol', '$')
        try:
            rows = self.ds.read_invoices()
            for row in reversed(rows):
                total_str = row.get('total', '0')
                try:
                    total_fmt = f"{sym}{float(total_str):.2f}"
                except ValueError:
                    total_fmt = total_str
                status = row.get('invoice_status', '').lower()
                if not status:
                    is_paid = row.get('paid', '').lower() in ('yes', 'true', '1')
                    status = 'paid' if is_paid else 'unpaid'
                if status == 'paid':
                    display = 'Paid'
                    tag = 'paid'
                elif status == 'unpaid':
                    display = 'Unpaid'
                    tag = 'unpaid'
                elif status == 'cancelled':
                    display = 'Cancelled'
                    tag = 'cancelled'
                elif status == 'void':
                    display = 'Void'
                    tag = 'void'
                else:
                    display = status.title()
                    tag = 'unpaid'
                self.hist_tree.insert('', 'end', tags=(tag,), values=(
                    row.get('invoice_number', ''),
                    storage_to_display(row.get('invoice_date', '')),
                    storage_to_display(row.get('due_date', '')),
                    row.get('client_name', ''),
                    total_fmt,
                    display
                ))
        except Exception:
            pass

    def _clients_tab_refresh(self):
        """Populate the Clients tab with per-client invoice stats."""
        for row in self.clients_tree.get_children():
            self.clients_tree.delete(row)
        sym = self.settings.get('currency_symbol', '$')
        stats = {}
        try:
            for row in self.ds.read_invoices():
                cname = row.get('client_name', '').strip()
                try:
                    total = float(row.get('total', '0'))
                except ValueError:
                    total = 0.0
                if cname:
                    cnt, ttl = stats.get(cname, (0, 0.0))
                    stats[cname] = (cnt + 1, ttl + total)
        except Exception:
            pass
        # Rows for known clients first
        shown = set()
        for c in self.clients:
            name = c['name']
            cnt, ttl = stats.get(name, (0, 0.0))
            self.clients_tree.insert('', 'end', values=(
                name, c.get('contact_name', ''), c.get('phone', ''),
                c.get('email', ''), cnt,
                f"{sym}{ttl:.2f}" if cnt else '-'
            ))
            shown.add(name)
        # Rows for clients only in invoices.csv (not in clients list)
        for name, (cnt, ttl) in sorted(stats.items()):
            if name not in shown:
                self.clients_tree.insert('', 'end', values=(
                    name, '', '', '', cnt, f"{sym}{ttl:.2f}"
                ))

    def _history_open_pdf(self):
        """Open the PDF for the selected history row."""
        sel = self.hist_tree.selection()
        if not sel:
            messagebox.showinfo("No selection", "Select an invoice from the list first.")
            return
        inv_num = self.hist_tree.item(sel[0], 'values')[0]
        pdf_path = self.ds.invoice_pdf_path(inv_num)
        if pdf_path.exists():
            _open_file(pdf_path)
        else:
            messagebox.showwarning("Not found", f"PDF not found:\n{pdf_path}")

    def _record_payment(self):
        """Open the payment recording dialog for the selected invoice."""
        sel = self.hist_tree.selection()
        if not sel:
            messagebox.showinfo("No selection", "Select an invoice from the list first.")
            return
        vals = self.hist_tree.item(sel[0], 'values')
        inv_num = vals[0]
        # Read current data for this invoice
        try:
            rows = self.ds.read_invoices()
            idx = next((i for i, r in enumerate(rows)
                        if r.get('invoice_number', '') == inv_num), None)
        except Exception as e:
            messagebox.showerror("Error", str(e))
            return
        if idx is None:
            messagebox.showwarning("Not found", f"Invoice {inv_num} not found in records.")
            return
        row = rows[idx]
        dlg = PaymentDialog(
            self.root,
            invoice_number=inv_num,
            client_name=row.get('client_name', ''),
            total=row.get('total', ''),
            paid=row.get('paid', '').lower() == 'yes',
            paid_date=row.get('paid_date', ''),
            payment_note=row.get('payment_note', ''),
        )
        if dlg.result is None:
            return
        # Write updated row back to CSV
        rows[idx]['paid']           = 'yes' if dlg.result['paid'] else ''
        rows[idx]['paid_date']      = dlg.result['paid_date']
        rows[idx]['payment_note']   = dlg.result['payment_note']
        rows[idx]['invoice_status'] = 'paid' if dlg.result['paid'] else 'unpaid'
        fieldnames = [
            'invoice_number', 'invoice_date', 'due_date', 'client_name',
            'client_address', 'notes', 'subtotal', 'gst', 'total',
            'paid', 'paid_date', 'payment_note', 'invoice_status', 'pdf_path'
        ]
        try:
            with open(self.invoices_csv_path, 'w', newline='', encoding='utf-8') as f:
                writer = csv.DictWriter(f, fieldnames=fieldnames, extrasaction='ignore')
                writer.writeheader()
                writer.writerows(rows)
        except Exception as e:
            messagebox.showerror("Save Error", str(e))
            return

        # If marked paid, update the matching ledger entry note or add a confirmed entry
        if dlg.result['paid']:
            inv_total = row.get('total', '')
            paid_date = dlg.result['paid_date']
            try:
                # Update any existing auto-posted ledger entry for this invoice
                ledger_rows = self.ds.read_ledger()
                updated = False
                for lr in ledger_rows:
                    if lr.get('reference') == inv_num and 'Auto-posted' in lr.get('notes', ''):
                        self.ds.update_ledger(lr['id'], {
                            'date':  paid_date or lr['date'],
                            'notes': f'Payment confirmed {storage_to_display(paid_date)}',
                        })
                        updated = True
                        break
                if not updated:
                    # No auto-posted entry exists — add a fresh payment entry
                    client_name = row.get('client_name', '')
                    self.ds.append_ledger({
                        'date':        paid_date or datetime.now().strftime('%Y-%m-%d'),
                        'type':        'in',
                        'category':    'Invoice Payment',
                        'description': f'Payment received — Invoice {inv_num} — {client_name}',
                        'amount':      inv_total,
                        'reference':   inv_num,
                        'notes':       f'Marked paid {storage_to_display(paid_date)}',
                    })
            except Exception:
                pass

        self._history_refresh()

    def _edit_invoice(self):
        """Open the selected invoice for editing via RecordMissingInvoiceDialog (reuse)."""
        sel = self.hist_tree.selection()
        if not sel:
            messagebox.showinfo('No selection', 'Select an invoice from the list first.')
            return
        inv_num = self.hist_tree.item(sel[0], 'values')[0]
        try:
            rows = self.ds.read_invoices()
            row  = next((r for r in rows if r.get('invoice_number', '') == inv_num), None)
        except Exception as e:
            messagebox.showerror('Error', str(e))
            return
        if row is None:
            messagebox.showwarning('Not found', f'Invoice {inv_num} not found in records.')
            return
        dlg = RecordMissingInvoiceDialog(
            self.root,
            ds=self.ds,
            settings=self.settings,
            currency_fn=lambda: self.settings.get('currency_symbol', '$'),
            clients=self.clients,
            prefill=row,
        )
        if dlg.result is None:
            return
        data = dlg.result
        rows_all = self.ds.read_invoices()
        idx = next((i for i, r in enumerate(rows_all)
                    if r.get('invoice_number') == data['invoice_number']), None)
        if idx is not None:
            rows_all[idx].update(data)
            self.ds._write_csv(self.ds.invoices_csv_path, INVOICE_FIELDS, rows_all)
            self.ds.audit('invoice_edited', f'#{data["invoice_number"]} {data["client_name"]}',
                          table='invoices', record_id=data['invoice_number'])
        else:
            self.ds.append_invoice(data)
        self._history_refresh()

    def _record_missing_invoice(self):
        """Open dialog to manually enter a past invoice into the system."""
        dlg = RecordMissingInvoiceDialog(
            self.root,
            ds=self.ds,
            settings=self.settings,
            currency_fn=lambda: self.settings.get('currency_symbol', '$'),
            clients=self.clients,
        )
        if dlg.result is None:
            return
        data = dlg.result
        inv_num = data['invoice_number']
        # If duplicate, update existing row; otherwise append
        rows = self.ds.read_invoices()
        existing_idx = next(
            (i for i, r in enumerate(rows) if r.get('invoice_number') == inv_num), None)
        if existing_idx is not None:
            rows[existing_idx].update(data)
            self.ds._write_csv(self.ds.invoices_csv_path, INVOICE_FIELDS, rows)
            self.ds.audit('invoice_missing_updated', f'#{inv_num} {data["client_name"]}',
                          table='invoices', record_id=inv_num)
        else:
            self.ds.append_invoice(data)
        messagebox.showinfo('Saved',
            f'Invoice #{inv_num} recorded for {data["client_name"]}.\n'
            'No PDF was generated.',
            parent=self.root)
        self._history_refresh()

    def _get_selected_invoice_number(self):
        """Return the invoice number of the selected history row, or None."""
        sel = self.hist_tree.selection()
        if not sel:
            messagebox.showinfo('No selection', 'Select an invoice from the list first.', parent=self.root)
            return None
        return self.hist_tree.item(sel[0], 'values')[0]

    def _cancel_invoice(self):
        """Mark the selected invoice as cancelled (keeps the record)."""
        inv_num = self._get_selected_invoice_number()
        if not inv_num:
            return
        if not messagebox.askyesno('Cancel Invoice',
                f'Cancel invoice #{inv_num}?\nThe record will remain but be marked Cancelled.',
                parent=self.root):
            return
        self.ds.update_invoice(inv_num, {
            'invoice_status': 'cancelled',
            'paid': '',
        })
        log_event('invoice', 'cancelled', f'#{inv_num}')
        self._history_refresh()

    def _reissue_invoice(self):
        """Regenerate the PDF for an existing invoice using the same invoice number."""
        inv_num = self._get_selected_invoice_number()
        if not inv_num:
            return
        rows = self.ds.read_invoices()
        row = next((r for r in rows if r.get('invoice_number') == inv_num), None)
        if not row:
            messagebox.showwarning('Not found', f'Invoice #{inv_num} not found.', parent=self.root)
            return
        pdf_path = self.ds.invoice_pdf_path(inv_num)
        if not messagebox.askyesno('Reissue Invoice',
                f'Reissue PDF for invoice #{inv_num}?\nPDF will be saved to:\n{pdf_path}',
                parent=self.root):
            return
        try:
            subtotal = float(row.get('subtotal', '0') or 0)
            gst = float(row.get('gst', '0') or 0)
            total = float(row.get('total', '0') or 0)
            items = [{
                'description': f"Invoice #{inv_num} — {row.get('client_name', '')}",
                'qty': 1, 'unit_price': subtotal, 'taxable': gst > 0,
                'subtotal': subtotal, 'gst': gst, 'total': total,
            }]
            self._create_pdf(
                pdf_path,
                invoice_number=inv_num,
                invoice_date=row.get('invoice_date', ''),
                due_date=row.get('due_date', ''),
                client_name=row.get('client_name', ''),
                client_address=row.get('client_address', ''),
                notes=row.get('notes', ''),
                items=items,
                subtotal=subtotal,
                gst=gst,
                total=total,
            )
            self.ds.update_invoice(inv_num, {'pdf_path': str(pdf_path)})
            messagebox.showinfo('Reissued', f'Invoice #{inv_num} PDF regenerated.', parent=self.root)
            _open_file(pdf_path)
        except Exception as e:
            messagebox.showerror('Reissue failed', str(e), parent=self.root)
        self._history_refresh()

    def _link_invoice_pdf(self):
        """Manually link an invoice to an external PDF file."""
        inv_num = self._get_selected_invoice_number()
        if not inv_num:
            return
        path = filedialog.askopenfilename(
            title=f'Link PDF for invoice #{inv_num}',
            defaultextension='.pdf',
            filetypes=[('PDF files', '*.pdf'), ('All files', '*.*')],
            parent=self.root)
        if not path:
            return
        self.ds.update_invoice(inv_num, {'pdf_path': str(Path(path))})
        log_event('invoice', 'pdf_linked', f'#{inv_num} -> {path}')
        messagebox.showinfo('Linked', f'Invoice #{inv_num} now links to:\n{path}', parent=self.root)
        self._history_refresh()

    def _open_help_guide(self):
        """Open the HTML help guide in the default browser.

        On first run the bundled help_guide.html is exported to the data
        directory so it persists even if the executable is moved.
        """
        import shutil
        import webbrowser

        # Preferred location: data directory
        guide = self.ds.data_dir / 'help_guide.html'

        if not guide.exists():
            # Find the bundled source
            if getattr(sys, 'frozen', False):
                bundle_dir = Path(sys._MEIPASS)
            else:
                bundle_dir = Path(__file__).parent
            src = bundle_dir / 'help_guide.html'
            if src.exists():
                try:
                    shutil.copy2(str(src), str(guide))
                except Exception:
                    guide = src  # fall back to reading from bundle

        if guide.exists():
            webbrowser.open(guide.as_uri())
        else:
            messagebox.showinfo('Help', 'Help guide file not found.\n'
                                'It will be regenerated on next launch.',
                                parent=self.root)

    def _reveal_invoice_folder(self):
        """Open the invoice PDF folder in Windows Explorer."""
        save_dir_setting = self.settings.get('pdf_save_dir', '').strip()
        folder = Path(save_dir_setting) if save_dir_setting else self.invoices_dir
        folder.mkdir(parents=True, exist_ok=True)
        _open_file(folder)

    # ------------------------------------------------------------------
    # Tools menu handlers
    # ------------------------------------------------------------------
    def _launch_receipt_generator(self):
        """Tools > Receipt Generator — launch the standalone receipt app."""
        import subprocess as _sp
        import sys as _sys

        # Determine the application directory (development or frozen)
        if getattr(_sys, 'frozen', False):
            app_dir = Path(_sys.executable).parent
        else:
            app_dir = Path(__file__).resolve().parent

        # Prefer the standalone EXE if it exists next to the main app
        exe_path = app_dir / 'Receipt Generator.exe'
        if exe_path.exists():
            _sp.Popen([str(exe_path)], shell=True)
            return

        # Otherwise run the Python script directly
        script_path = app_dir / 'receipt_generator' / 'receipt_app.py'
        if script_path.exists():
            _sp.Popen([_sys.executable, str(script_path)], shell=True)
            return

        messagebox.showinfo(
            'Receipt Generator not found',
            'Receipt Generator could not be located.\n'
            'Expected one of:\n'
            f'  {exe_path}\n'
            f'  {script_path}',
            parent=self.root
        )

    def _new_invoice(self):
        """Clear the invoice form (Edit > New Invoice)."""
        self.notebook.select(0)
        self._clear_form(reset_invoice_number=False)

    def _add_ledger_menu(self):
        """Edit > Add Ledger Entry — switch to Ledger tab."""
        self.notebook.select(self._ledger_tab.frame)

    def _open_users(self):
        """Users > Manage Users — open the user management dialog."""
        UsersDialog(self.root, self.ds)

    def _open_settings(self):
        """Open the settings dialog and apply changes."""
        # Pass current data_dir so Config tab shows correct value
        settings_with_cfg = dict(self.settings)
        settings_with_cfg['_data_dir'] = self.ds.config.get('data_dir', '')
        dlg = SettingsDialog(self.root, settings_with_cfg,
                              backup_mgr=getattr(self, '_backup_mgr', None))
        if dlg.result is not None:
            new_data_dir = dlg.result.pop('_data_dir', '')
            if new_data_dir != self.ds.config.get('data_dir', ''):
                self.ds.update_data_dir(new_data_dir)
                setup_logging(self.ds.data_dir)   # move log file to new data dir
                self.settings_path      = self.ds.settings_path
                self.service_items_path = self.ds.service_items_path
                self.clients_path       = self.ds.clients_path
                self.invoices_dir       = self.ds.invoices_dir
                self.invoices_csv_path  = self.ds.invoices_csv_path
                self.ds.ensure_files()
                messagebox.showinfo('Data directory changed',
                    'Data folder updated. Restart to ensure all tabs use the new location.',
                    parent=self.root)
            self.settings.update(dlg.result)
            self._save_settings()
            self._populate_invoice_details()
            self._backup_mgr.apply_new_settings()
            messagebox.showinfo('Settings', 'Settings saved. Some changes take effect on next invoice.')

    def _reload_all(self):
        """Re-read every CSV and settings.json from disk and refresh all tabs."""
        self.settings      = self._load_settings()
        self.service_items = self._load_service_items()
        self.clients       = self._load_clients()
        self._refresh_client_combo()
        self._populate_invoice_details()
        self._history_refresh()
        self._clients_tab_refresh()
        for tab_obj in (
            getattr(self, '_ledger_tab',   None),
            getattr(self, '_reports_tab',  None),
        ):
            if tab_obj and hasattr(tab_obj, 'refresh'):
                try:
                    tab_obj.refresh()
                except Exception:
                    pass

    def _on_close(self):
        """Gracefully shut down — run on-exit backup then destroy."""
        try:
            self._backup_mgr.stop()
            self._backup_mgr.run_now(on_exit=True)
        except Exception:
            pass
        self.root.destroy()

    def _backup_now(self):
        """Tools > Backup Now — run an immediate silent backup."""
        try:
            self._backup_mgr.run_now(on_exit=False)
            bdir = self.settings.get('backup_dir', '').strip() or str(self.ds.data_dir / 'backups')
            messagebox.showinfo('Backup complete',
                f'Backup written to:\n{bdir}', parent=self.root)
        except Exception as e:
            messagebox.showerror('Backup failed', str(e), parent=self.root)

    def _export_data(self):
        """Export all data files to a zip archive."""
        path = filedialog.asksaveasfilename(
            title='Export All Data',
            initialfile=f'invoicer_backup_{datetime.now().strftime("%Y%m%d_%H%M%S")}.zip',
            defaultextension='.zip',
            filetypes=[('Zip archive', '*.zip'), ('All files', '*.*')],
            parent=self.root)
        if not path:
            return
        try:
            self.ds.export_all(Path(path))
            # Read manifest to show a detailed summary
            import zipfile as _zf, json as _json
            with _zf.ZipFile(path, 'r') as z:
                manifest = _json.loads(z.read('manifest.json'))
            files = manifest.get('files', [])
            data_files = [f for f in files if f.startswith('data/')]
            pdf_files  = [f for f in files if f.startswith('invoices/')]
            log_files  = [f for f in files if f.startswith('logs/')]
            messagebox.showinfo(
                'Export complete',
                f'Backup saved to:\n{path}\n\n'
                f'  Data files:    {len(data_files)}\n'
                f'  Invoice PDFs:  {len(pdf_files)}\n'
                f'  Log files:     {len(log_files)}\n'
                f'  Total:         {len(files)} files\n\n'
                f'Exported at: {manifest.get("exported_at", "")}\n'
                f'App version: {manifest.get("app_version", "")}',
                parent=self.root)
        except Exception as e:
            messagebox.showerror('Export failed', str(e), parent=self.root)

    def _import_data(self):
        """Import data from a previously exported zip archive."""
        path = filedialog.askopenfilename(
            title='Import Data',
            filetypes=[('Zip archive', '*.zip'), ('All files', '*.*')],
            parent=self.root)
        if not path:
            return

        # Show manifest info before confirming
        try:
            import zipfile as _zf, json as _json
            with _zf.ZipFile(path, 'r') as z:
                if 'manifest.json' in z.namelist():
                    m = _json.loads(z.read('manifest.json'))
                    files = m.get('files', [])
                    info = (
                        f'Backup exported: {m.get("exported_at","unknown")}\n'
                        f'App version:     {m.get("app_version","unknown")}\n'
                        f'Data files:      {len([f for f in files if f.startswith("data/")])}\n'
                        f'Invoice PDFs:    {len([f for f in files if f.startswith("invoices/")])}\n'
                        f'Log files:       {len([f for f in files if f.startswith("logs/")])}\n'
                    )
                else:
                    info = '(Legacy backup — no manifest)\n'
        except Exception:
            info = ''

        overwrite = messagebox.askyesno(
            'Overwrite existing?',
            f'{info}\nOverwrite existing data files?\n\nChoose No to only import files that don\'t already exist.',
            parent=self.root)
        try:
            results = self.ds.import_all(Path(path), overwrite=overwrite)
            imported = [n for n, a in results if a == 'imported']
            skipped  = [n for n, a in results if a == 'skipped']
            messagebox.showinfo('Import complete',
                f'Import finished.\n\n'
                f'  Imported: {len(imported)} files\n'
                f'  Skipped:  {len(skipped)} files (already existed)',
                parent=self.root)
            # Reload all in-memory data
            self.settings = self._load_settings()
            self.service_items = self._load_service_items()
            self.clients = self._load_clients()
            self._refresh_client_combo()
            self._populate_invoice_details()
            self._history_refresh()
            # Trigger other tabs to reload on next switch
            for tab_obj in (
                getattr(self, '_ledger_tab',   None),
                getattr(self, '_reports_tab',  None),
            ):
                if tab_obj and hasattr(tab_obj, 'refresh'):
                    try:
                        tab_obj.refresh()
                    except Exception:
                        pass
        except Exception as e:
            messagebox.showerror('Import failed', str(e), parent=self.root)

    def _migrate_from_v15(self):
        """
        Import data from a V1.5 (or earlier) data folder or zip, then
        auto-migrate all CSV schemas to the current version.
        Supports both plain folders and .zip archives.
        """
        path = filedialog.askopenfilename(
            title='Select V1.5 data folder or zip archive',
            filetypes=[
                ('Zip archive', '*.zip'),
                ('All files', '*.*'),
            ],
            parent=self.root)
        # If user cancels file dialog, try askdirectory for folder import
        if not path:
            path = filedialog.askdirectory(
                title='Or choose the V1.5 data folder',
                parent=self.root)
        if not path:
            return

        src = Path(path)
        overwrite = messagebox.askyesno(
            'Overwrite existing data?',
            'Overwrite any existing data files with V1.5 data?\n\n'
            'Choose No to skip files that already exist in the current data folder.\n\n'
            'Tip: Export a backup first (Tools → Export All Data).',
            parent=self.root)

        try:
            if src.is_dir():
                result = self.ds.import_from_folder(src, overwrite=overwrite)
                copied   = result['copied']
                skipped  = result['skipped']
                migrated = result['migrated']
            else:
                # zip path — use existing import_all then report migration separately
                raw = self.ds.import_all(src, overwrite=overwrite)
                copied   = [n for n, a in raw if a == 'imported']
                skipped  = [n for n, a in raw if a == 'skipped']
                migrated = self.ds.migrate_all()
        except Exception as e:
            messagebox.showerror('Migration failed', str(e), parent=self.root)
            return

        # Build human-readable report
        lines = []
        lines.append(f'Files copied:  {len(copied)}')
        if skipped:
            lines.append(f'Files skipped: {len(skipped)}  (already existed)')
        lines.append('')
        if migrated:
            lines.append('Schema upgrades applied:')
            for fname, cols in migrated.items():
                lines.append(f'  {fname}: added {", ".join(cols)}')
        else:
            lines.append('All schemas already up-to-date — no migration needed.')

        messagebox.showinfo(
            'Migration complete',
            '\n'.join(lines),
            parent=self.root)

        # Reload all tabs
        self.settings = self._load_settings()
        self.service_items = self._load_service_items()
        self.clients = self._load_clients()
        self._refresh_client_combo()
        self._populate_invoice_details()
        self._history_refresh()
        for tab_obj in (
            getattr(self, '_ledger_tab', None),
            getattr(self, '_reports_tab', None),
        ):
            if tab_obj and hasattr(tab_obj, 'refresh'):
                try:
                    tab_obj.refresh()
                except Exception:
                    pass

    def _export_history_csv(self):
        """Export the invoice history list to a CSV file."""
        path = filedialog.asksaveasfilename(
            title='Export Invoice History',
            initialfile=f'invoice_history_{datetime.now().strftime("%Y%m%d")}.csv',
            defaultextension='.csv',
            filetypes=[('CSV files', '*.csv'), ('All files', '*.*')],
            parent=self.root)
        if not path:
            return
        try:
            rows = self.ds.read_invoices()
            import csv as _csv
            with open(path, 'w', newline='', encoding='utf-8') as f:
                if rows:
                    w = _csv.DictWriter(f, fieldnames=rows[0].keys())
                    w.writeheader()
                    w.writerows(rows)
            messagebox.showinfo('Exported',
                f'{len(rows)} invoices exported to:\n{path}', parent=self.root)
        except Exception as e:
            messagebox.showerror('Export failed', str(e), parent=self.root)

    def _open_clients(self):
        """Open the client management dialog."""
        dlg = ClientsDialog(self.root, self.clients, self.clients_path)
        if dlg.changed:
            self.clients = self._load_clients()
            self._refresh_client_combo()

    def _refresh_client_combo(self):
        """Rebuild the client picker dropdown values."""
        names = [''] + [c['name'] for c in self.clients]
        self.client_combo['values'] = names

    def _client_selected(self, event=None):
        """Auto-fill client name and address when a saved client is chosen."""
        selected = self.client_pick_var.get()
        if not selected:
            return
        for c in self.clients:
            if c['name'] == selected:
                self.client_name_var.set(c['name'])
                self.client_address_text.delete('1.0', 'end')
                if c.get('address'):
                    self.client_address_text.insert('1.0', c['address'])
                break

    def _open_catalogue(self):
        """Open the service catalogue management dialog."""
        dlg = CatalogueDialog(self.root, self.service_items, self.service_items_path)
        if dlg.changed:
            self.service_items = self._load_service_items()
            service_options = [''] + [item['description'] for item in self.service_items]
            self.service_combo['values'] = service_options


def main():
    # ------------------------------------------------------------------
    # Bootstrap logging before anything else
    # ------------------------------------------------------------------
    _base = Path(sys.executable).parent if getattr(sys, 'frozen', False) else Path(__file__).parent
    setup_logging(_base)          # initial log in app folder; re-pointed after DS loads

    # Delayed-run flag (set when launched from Windows startup)
    if '--delayed' in sys.argv:
        import time
        time.sleep(15)

    root = tk.Tk()
    root.withdraw()
    apply_theme(root)

    # Capture Tkinter internal errors (callbacks etc.) into the log
    def _tk_error(exc, val, tb):
        _log.error('Tkinter callback error:\n%s',
                   ''.join(traceback.format_exception(exc, val, tb)))
    root.report_callback_exception = _tk_error

    # Initialise DataStore and show login before building the main UI
    ds = DataStore(_base)
    ds.ensure_files()
    ds.migrate_all()
    login = LoginDialog(root, ds)
    if not login.result:
        return
    root.deiconify()

    app = InvoiceApp(root, ds=ds, current_user=login.result.get('username', 'admin'))

    # Re-point log to data dir once DataStore is initialised
    try:
        setup_logging(app.ds.data_dir)
    except Exception:
        pass

    _log.info('UI ready — entering mainloop')
    try:
        root.mainloop()
    finally:
        _log.info('=== Invoice Generator exited ===')
        logging.shutdown()


if __name__ == '__main__':
    main()