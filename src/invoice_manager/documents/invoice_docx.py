"""Generate a Microsoft Word .docx version of an invoice."""

from __future__ import annotations

from decimal import Decimal
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from invoice_manager.domain.money import Money
from invoice_manager.persistence.models import Invoice


def _get(settings: dict[str, Any], key: str, default: str = "") -> str:
    value = settings.get(key, default)
    if not value:
        return default
    return str(value)


def _fmt(cents: int) -> str:
    return Money(cents=cents).__str__()


def generate_invoice_docx(invoice: Invoice, settings: dict[str, Any], output_path: Path) -> Path:
    """Create a .docx document for the invoice."""
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    document = Document()

    business = _get(settings, "business_name", "Invoice")
    para = document.add_paragraph()
    run = para.add_run(business)
    run.bold = True
    run.font.size = para.runs[0].font.size
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    address = _get(settings, "business_address")
    if address:
        document.add_paragraph(address)

    gst_rate = Decimal(_get(settings, "gst_rate", "0.0") or "0.0")
    doc_title = (
        _get(settings, "invoice_title_tax", "TAX INVOICE")
        if gst_rate > 0
        else _get(settings, "invoice_title", "INVOICE")
    )
    title = document.add_paragraph()
    run = title.add_run(f"{doc_title} — {invoice.number}")
    run.bold = True
    run.font.size = title.runs[0].font.size

    document.add_paragraph()
    meta = [
        (_get(settings, "invoice_date_label", "Date:"), str(invoice.issue_date)),
        (_get(settings, "invoice_due_date_label", "Due date:"), str(invoice.due_date or "")),
        (_get(settings, "invoice_client_label", "Client:"), invoice.client_name),
        (_get(settings, "invoice_address_label", "Address:"), invoice.client_address or ""),
    ]
    for label, value in meta:
        p = document.add_paragraph()
        p.add_run(label).bold = True
        p.add_run(f" {value}")

    table = document.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    headers = [
        _get(settings, "invoice_description_header", "Description"),
        _get(settings, "invoice_qty_header", "Qty"),
        _get(settings, "invoice_unit_header", "Unit"),
        _get(settings, "invoice_price_header", "Price"),
        _get(settings, "invoice_gst_header", "GST"),
        _get(settings, "invoice_total_header", "Total"),
    ]
    for idx, header in enumerate(headers):
        hdr_cells[idx].text = header
        for paragraph in hdr_cells[idx].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    for item in invoice.items:
        row_cells = table.add_row().cells
        row_cells[0].text = item.description
        row_cells[1].text = str(item.quantity)
        row_cells[2].text = item.unit or "ea"
        row_cells[3].text = _fmt(item.unit_price_cents)
        row_cells[4].text = _fmt(item.gst_cents)
        row_cells[5].text = _fmt(item.total_cents)

    for label, amount in [
        (_get(settings, "invoice_subtotal_label", "Subtotal"), _fmt(invoice.subtotal_cents)),
        (_get(settings, "invoice_gst_label", "GST"), _fmt(invoice.gst_cents)),
        (_get(settings, "invoice_total_label", "Total"), _fmt(invoice.total_cents)),
    ]:
        row_cells = table.add_row().cells
        row_cells[0].text = ""
        row_cells[1].text = ""
        row_cells[2].text = ""
        row_cells[3].text = label
        row_cells[4].text = ""
        row_cells[5].text = amount
        for paragraph in row_cells[3].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        for paragraph in row_cells[5].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    document.add_paragraph()

    bank_name = _get(settings, "bank_name")
    bsb = _get(settings, "bank_bsb")
    account = _get(settings, "bank_account")
    account_name = _get(settings, "bank_account_name")
    if bank_name or account:
        p = document.add_paragraph()
        p.add_run(_get(settings, "invoice_payment_details_label", "Payment details")).bold = True
        payment_line = (
            f"{_get(settings, 'invoice_bank_label', 'Bank:')} {bank_name}  |  "
            f"{_get(settings, 'invoice_bsb_label', 'BSB:')} {bsb}  |  "
            f"{_get(settings, 'invoice_account_label', 'Account:')} {account}  |  "
            f"{_get(settings, 'invoice_account_name_label', 'Name:')} {account_name}"
        )
        document.add_paragraph(payment_line)

    if invoice.notes:
        p = document.add_paragraph()
        p.add_run(_get(settings, "invoice_notes_label", "Notes:")).bold = True
        p.add_run(f" {invoice.notes}")

    thank_you = _get(settings, "invoice_thank_you", "Thank you for your business!")
    if thank_you:
        document.add_paragraph(thank_you)

    document.save(str(output_path))
    return output_path
