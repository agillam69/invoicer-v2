"""Generate receipt documents in Microsoft Word format."""

from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document

from invoice_manager.domain.money import Money
from invoice_manager.persistence.models import Invoice, Payment, Receipt


def generate_receipt_docx(
    receipt: Payment | Receipt,
    invoice: Invoice | None,
    settings: dict[str, Any],
    output_path: Path,
) -> Path:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.add_heading(str(settings.get("business_name") or "Receipt"), 0)
    address = str(settings.get("business_address") or "")
    if address:
        document.add_paragraph(address)
    number = receipt.receipt_number if isinstance(receipt, Payment) else receipt.number
    document.add_heading(f"{settings.get('receipt_title') or 'RECEIPT'} — {number}", 1)
    table = document.add_table(rows=0, cols=2)
    if invoice is not None:
        paid = sum(p.amount_cents for p in invoice.payments if not p.is_reversed)
        credits = sum(c.amount_cents for c in invoice.credits)
        rows = [
            ("Received from", invoice.client_name),
            ("Invoice", invoice.number),
            ("Invoice date", str(invoice.issue_date)),
            ("Invoice amount", str(Money(cents=invoice.total_cents))),
            ("Amount paid", str(Money(cents=receipt.amount_cents))),
            ("Payment method", receipt.method or ""),
            ("Reference", receipt.reference or ""),
            ("Amount outstanding", str(Money(cents=max(0, invoice.total_cents - paid - credits)))),
        ]
    else:
        assert isinstance(receipt, Receipt)
        rows = [
            ("Received from", receipt.client_name),
            ("Date", str(receipt.date)),
            ("Amount received", str(Money(cents=receipt.amount_cents))),
            ("Payment method", receipt.method or ""),
            ("Reference", receipt.reference or ""),
            ("For", receipt.description or ""),
            ("Notes", receipt.notes or ""),
        ]
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
    thank_you = str(settings.get("receipt_thank_you") or "Thank you for your payment.")
    if thank_you:
        document.add_paragraph(thank_you)
    document.save(str(output_path))
    return output_path
