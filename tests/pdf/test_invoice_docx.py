from datetime import date
from decimal import Decimal

from docx import Document

from invoice_manager.documents.invoice_docx import generate_invoice_docx
from invoice_manager.persistence.models import Client, Invoice, InvoiceItem


def test_invoice_docx_contains_number_and_total(tmp_path):
    client = Client(name="Acme Corp", address="123 Main St")
    invoice = Invoice(
        number="INV-0001",
        sequence_number=1,
        issue_date=date(2026, 1, 15),
        due_date=date(2026, 2, 15),
        client=client,
        client_name=client.name,
        client_address=client.address,
        subtotal_cents=10000,
        gst_cents=1000,
        total_cents=11000,
        is_draft=False,
    )
    invoice.items.append(
        InvoiceItem(
            invoice=invoice,
            description="Consulting",
            quantity=1,
            unit="ea",
            unit_price_cents=10000,
            taxable=True,
            subtotal_cents=10000,
            gst_cents=1000,
            total_cents=11000,
            sort_order=0,
        )
    )
    settings = {
        "business_name": "Test Business",
        "business_address": "456 Market St",
        "gst_rate": Decimal("0.10"),
        "bank_name": "Test Bank",
        "bank_bsb": "000-000",
        "bank_account": "12345678",
        "bank_account_name": "Test Business",
        "thank_you_note": "Thanks!",
    }
    output = tmp_path / "invoice.docx"
    generate_invoice_docx(invoice, settings, output)

    assert output.exists()
    doc = Document(output)
    text = " ".join(paragraph.text for paragraph in doc.paragraphs)
    for table in doc.tables:
        for table_row in table.rows:
            text += " " + " ".join(cell.text for cell in table_row.cells)
    assert "INV-0001" in text
    assert "Acme Corp" in text
    assert "$110.00" in text
    assert "Consulting" in text
